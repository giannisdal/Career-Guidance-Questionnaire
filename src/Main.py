import sys
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import Qt
import pandas as pd
from PyQt5.QtWidgets import *
from PyQt5.QtCore import Qt
from PyQt5.QtGui import *
import sys
from PyQt5 import uic
from docx import Document
import matplotlib.pyplot as plt
import numpy as np
# from mplwidget import MplWidget
import ctypes

myappid = 'mycompany.myproduct.subproduct.version'  # arbitrary string
ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)


class TableModel(QtCore.QAbstractTableModel):

    def __init__(self, data):
        super(TableModel, self).__init__()
        self._data = data

    def data(self, index, role):
        if role == Qt.DisplayRole:
            value = self._data.iloc[index.row(), index.column()]
            return str(value)
        elif role == QtCore.Qt.TextAlignmentRole:
            return QtCore.Qt.AlignCenter

    def rowCount(self, index):
        return self._data.shape[0]

    def columnCount(self, index):
        return self._data.shape[1]

    def headerData(self, section, orientation, role):
        # section is the index of the column/row.
        if role == Qt.DisplayRole:
            if orientation == Qt.Horizontal:
                return str(self._data.columns[section])

            if orientation == Qt.Vertical:
                return str(self._data.index[section])


class MainWindow(QtWidgets.QMainWindow):

    def __init__(self):
        super().__init__()

        plt.rc('font', weight='bold')
        plt.rcParams["font.family"] = "serif"
        plt.rcParams["mathtext.fontset"] = "dejavuserif"
        # plt.style.use('ggplot')

        uic.loadUi(r"C:\Users\Public\Documents\Python\Ui2.ui", self)
        self.initUI()

    def initUI(self):
        self.setWindowTitle("Επαγγελματικός Προσανατολισμός")
        self.setWindowIcon(QtGui.QIcon(r"C:\Users\Public\Documents\Python\earth.png"))
        self.document = Document(r"C:\Users\Public\Documents\Python\aksies.docx")
        self.second_document = Document(r"C:\Users\Public\Documents\Python\autopepoithisi.docx")
        self.holland_document = Document(r"C:\Users\Public\Documents\Python\holland.docx")
        self.doc1 = Document(r"C:\Users\Public\Documents\Python\doc1.docx")
        self.data()
        self.second_data()
        self.holland_data()
        self.doc()
        self.table1()
        self.table2()
        self.table3()
        self.table4_1()
        self.table4_2()
        self.table4_3()
        self.table4_4()
        self.table4_5()
        self.table4_6()

        self.table4_7()
        self.table4_8()
        self.table4_9()
        self.table4_10()
        self.table4_11()
        self.table4_12()

        self.table4_13()
        self.table4_14()
        self.table4_15()
        self.table4_16()
        self.table4_17()
        self.table4_18()

        self.table4_19()
        self.table4_20()

        self.table5_1()
        self.table5_2()
        self.table5_3()
        self.table5_4()
        self.table5_5()
        self.table5_6()

        self.pushButton.clicked.connect(self.checkboxChanged)
        self.pushButton.clicked.connect(self.update_graph_1)
        self.pushButton.clicked.connect(self.update_graph_1_2)
        self.pushButton.clicked.connect(self.pressed_button1)

        self.pushButton_2.clicked.connect(self.checkboxChanged_2)
        self.pushButton_2.clicked.connect(self.table1_results)
        self.pushButton_2.clicked.connect(self.update_graph_2)
        self.pushButton_2.clicked.connect(self.pressed_button2)

        self.pushButton_3.clicked.connect(self.checkboxChanged_3)
        self.pushButton_3.clicked.connect(self.table1_results_2)
        self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_3.clicked.connect(self.update_graph_3_1)
        self.pushButton_3.clicked.connect(self.pressed_button3)

        self.pushButton_7.clicked.connect(self.checkboxChanged_4_1)
        self.pushButton_7.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_7.clicked.connect(self.pressed_button4)

        self.pushButton_8.clicked.connect(self.checkboxChanged_4_2)
        self.pushButton_8.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_8.clicked.connect(self.pressed_button5)

        self.pushButton_9.clicked.connect(self.checkboxChanged_4_3)
        self.pushButton_9.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_9.clicked.connect(self.pressed_button6)

        self.pushButton_10.clicked.connect(self.checkboxChanged_4_4)
        self.pushButton_10.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_10.clicked.connect(self.pressed_button7)

        self.pushButton_11.clicked.connect(self.checkboxChanged_4_5)
        self.pushButton_11.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_11.clicked.connect(self.pressed_button8)

        self.pushButton_12.clicked.connect(self.checkboxChanged_4_6)
        self.pushButton_12.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_12.clicked.connect(self.pressed_button9)

        self.pushButton_13.clicked.connect(self.checkboxChanged_4_7)
        self.pushButton_13.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_13.clicked.connect(self.pressed_button10)

        self.pushButton_14.clicked.connect(self.checkboxChanged_4_8)
        self.pushButton_14.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_14.clicked.connect(self.pressed_button11)

        self.pushButton_15.clicked.connect(self.checkboxChanged_4_9)
        self.pushButton_15.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_15.clicked.connect(self.pressed_button12)

        self.pushButton_16.clicked.connect(self.checkboxChanged_4_10)
        self.pushButton_16.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_16.clicked.connect(self.pressed_button13)

        self.pushButton_17.clicked.connect(self.checkboxChanged_4_11)
        self.pushButton_17.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_17.clicked.connect(self.pressed_button14)

        self.pushButton_18.clicked.connect(self.checkboxChanged_4_12)
        self.pushButton_18.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_18.clicked.connect(self.pressed_button15)

        self.pushButton_19.clicked.connect(self.checkboxChanged_4_13)
        self.pushButton_19.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_19.clicked.connect(self.pressed_button16)

        self.pushButton_20.clicked.connect(self.checkboxChanged_4_14)
        self.pushButton_20.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_20.clicked.connect(self.pressed_button17)

        self.pushButton_21.clicked.connect(self.checkboxChanged_4_15)
        self.pushButton_21.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_21.clicked.connect(self.pressed_button18)

        self.pushButton_22.clicked.connect(self.checkboxChanged_4_16)
        self.pushButton_22.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_22.clicked.connect(self.pressed_button19)

        self.pushButton_23.clicked.connect(self.checkboxChanged_4_17)
        self.pushButton_23.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_23.clicked.connect(self.pressed_button20)

        self.pushButton_24.clicked.connect(self.checkboxChanged_4_18)
        self.pushButton_24.clicked.connect(self.hol_results)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_24.clicked.connect(self.pressed_button21)

        self.pushButton_25.clicked.connect(self.checkboxChanged_4_19)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        self.pushButton_25.clicked.connect(self.checkboxChanged_4_20)
        self.pushButton_25.clicked.connect(self.R)
        self.pushButton_25.clicked.connect(self.I)
        self.pushButton_25.clicked.connect(self.A)
        self.pushButton_25.clicked.connect(self.S)
        self.pushButton_25.clicked.connect(self.E)
        self.pushButton_25.clicked.connect(self.C)
        self.pushButton_25.clicked.connect(self.hol_results)
        self.pushButton_25.clicked.connect(self.pressed_button22)

        self.pushButton_4.clicked.connect(self.stoixeia)
        self.pushButton_4.clicked.connect(self.word)
        # self.pushButton_4.clicked.connect(self.pressed_button22)
        # self.pushButton_3.clicked.connect(self.update_graph_3)
        # self.pushButton_25.clicked.connect(self.pressed_button22)

        self.labelResult1_1 = QLabel(self)
        self.labelResult1_1.resize(0, 0)
        self.labelResult2_1 = QLabel(self)
        self.labelResult2_1.resize(0, 0)
        self.labelResult3_1 = QLabel(self)
        self.labelResult3_1.resize(0, 0)

        self.labelResult0_2 = QLabel(self)
        self.labelResult0_2.resize(0, 0)
        self.labelResult1_2 = QLabel(self)
        self.labelResult1_2.resize(0, 0)
        self.labelResult2_2 = QLabel(self)
        self.labelResult2_2.resize(0, 0)
        self.labelResult3_2 = QLabel(self)
        self.labelResult3_2.resize(0, 0)

        self.labelResult1_3 = QLabel(self)
        self.labelResult1_3.resize(0, 0)
        self.labelResult2_3 = QLabel(self)
        self.labelResult2_3.resize(0, 0)
        self.labelResult3_3 = QLabel(self)
        self.labelResult3_3.resize(0, 0)
        self.labelResult4_3 = QLabel(self)
        self.labelResult4_3.resize(0, 0)
        self.labelResult5_3 = QLabel(self)
        self.labelResult5_3.resize(0, 0)

        self.labelResult1_4_1 = QLabel(self)
        self.labelResult1_4_1.resize(0, 0)
        self.labelResult2_4_1 = QLabel(self)
        self.labelResult2_4_1.resize(0, 0)

        self.labelResult1_4_2 = QLabel(self)
        self.labelResult1_4_2.resize(0, 0)
        self.labelResult2_4_2 = QLabel(self)
        self.labelResult2_4_2.resize(0, 0)

        self.labelResult1_4_3 = QLabel(self)
        self.labelResult1_4_3.resize(0, 0)
        self.labelResult2_4_3 = QLabel(self)
        self.labelResult2_4_3.resize(0, 0)

        self.labelResult1_4_4 = QLabel(self)
        self.labelResult1_4_4.resize(0, 0)
        self.labelResult2_4_4 = QLabel(self)
        self.labelResult2_4_4.resize(0, 0)

        self.labelResult1_4_5 = QLabel(self)
        self.labelResult1_4_5.resize(0, 0)
        self.labelResult2_4_5 = QLabel(self)
        self.labelResult2_4_5.resize(0, 0)

        self.labelResult1_4_6 = QLabel(self)
        self.labelResult1_4_6.resize(0, 0)
        self.labelResult2_4_6 = QLabel(self)
        self.labelResult2_4_6.resize(0, 0)

        self.labelResult1_4_7 = QLabel(self)
        self.labelResult1_4_7.resize(0, 0)
        self.labelResult2_4_7 = QLabel(self)
        self.labelResult2_4_7.resize(0, 0)

        self.labelResult1_4_8 = QLabel(self)
        self.labelResult1_4_8.resize(0, 0)
        self.labelResult2_4_8 = QLabel(self)
        self.labelResult2_4_8.resize(0, 0)

        self.labelResult1_4_9 = QLabel(self)
        self.labelResult1_4_9.resize(0, 0)
        self.labelResult2_4_9 = QLabel(self)
        self.labelResult2_4_9.resize(0, 0)

        self.labelResult1_4_10 = QLabel(self)
        self.labelResult1_4_10.resize(0, 0)
        self.labelResult2_4_10 = QLabel(self)
        self.labelResult2_4_10.resize(0, 0)

        self.labelResult1_4_11 = QLabel(self)
        self.labelResult1_4_11.resize(0, 0)
        self.labelResult2_4_11 = QLabel(self)
        self.labelResult2_4_11.resize(0, 0)

        self.labelResult1_4_12 = QLabel(self)
        self.labelResult1_4_12.resize(0, 0)
        self.labelResult2_4_12 = QLabel(self)
        self.labelResult2_4_12.resize(0, 0)

        self.labelResult1_4_13 = QLabel(self)
        self.labelResult1_4_13.resize(0, 0)
        self.labelResult2_4_13 = QLabel(self)
        self.labelResult2_4_13.resize(0, 0)

        self.labelResult1_4_14 = QLabel(self)
        self.labelResult1_4_14.resize(0, 0)
        self.labelResult2_4_14 = QLabel(self)
        self.labelResult2_4_14.resize(0, 0)

        self.labelResult1_4_15 = QLabel(self)
        self.labelResult1_4_15.resize(0, 0)
        self.labelResult2_4_15 = QLabel(self)
        self.labelResult2_4_15.resize(0, 0)

        self.labelResult1_4_16 = QLabel(self)
        self.labelResult1_4_16.resize(0, 0)
        self.labelResult2_4_16 = QLabel(self)
        self.labelResult2_4_16.resize(0, 0)

        self.labelResult1_4_17 = QLabel(self)
        self.labelResult1_4_17.resize(0, 0)
        self.labelResult2_4_17 = QLabel(self)
        self.labelResult2_4_17.resize(0, 0)

        self.labelResult1_4_18 = QLabel(self)
        self.labelResult1_4_18.resize(0, 0)
        self.labelResult2_4_18 = QLabel(self)
        self.labelResult2_4_18.resize(0, 0)

        self.labelResult1_4_19 = QLabel(self)
        self.labelResult1_4_19.resize(0, 0)
        self.labelResult2_4_19 = QLabel(self)
        self.labelResult2_4_19.resize(0, 0)
        self.labelResult3_4_19 = QLabel(self)
        self.labelResult3_4_19.resize(0, 0)
        self.labelResult4_4_19 = QLabel(self)
        self.labelResult4_4_19.resize(0, 0)
        self.labelResult5_4_19 = QLabel(self)
        self.labelResult5_4_19.resize(0, 0)
        self.labelResult6_4_19 = QLabel(self)
        self.labelResult6_4_19.resize(0, 0)

        self.labelResult1_4_20 = QLabel(self)
        self.labelResult1_4_20.resize(0, 0)
        self.labelResult2_4_20 = QLabel(self)
        self.labelResult2_4_20.resize(0, 0)
        self.labelResult3_4_20 = QLabel(self)
        self.labelResult3_4_20.resize(0, 0)
        self.labelResult4_4_20 = QLabel(self)
        self.labelResult4_4_20.resize(0, 0)
        self.labelResult5_4_20 = QLabel(self)
        self.labelResult5_4_20.resize(0, 0)
        self.labelResult6_4_20 = QLabel(self)
        self.labelResult6_4_20.resize(0, 0)

        self.lineEdit_2.text()

        self.table1_results()
        # self.table1_results_2()

        self.tabWidget_A.tabBar().setTabTextColor(0, QColor("dimgrey"))
        for i in range(0, 3):
            self.tabWidget.tabBar().setTabTextColor(i, QColor("dimgrey"))

        self.tabWidget_A.tabBar().setTabTextColor(1, QColor("darkmagenta"))
        for i in range(0, 2):
            self.tabWidget_2.tabBar().setTabTextColor(i, QColor("darkmagenta"))
        self.tabWidget_A.tabBar().setTabTextColor(2, QColor("mediumblue"))
        for i in range(0, 5):
            self.tabWidget_3.tabBar().setTabTextColor(i, QColor("mediumblue"))
        self.show()

    def stoixeia(self):
        dframe = pd.DataFrame({"name": [self.lineEdit.text()], "age": [self.lineEdit_2.text()],
                               "sex": [self.comboBox.currentText()], "learn": [self.comboBox_2.currentText()],
                               "job": [self.lineEdit_3.text()]})
        print(dframe)
        if ((self.comboBox.currentText() == "") or (self.comboBox_2.currentText() == "")
                or (self.lineEdit.text() == "") or (self.lineEdit_2.text() == "") or (self.lineEdit_3.text() == "")):
            self.label_49.setText(u"Σφάλμα Συμπλήρωσης")
        elif not self.lineEdit_2.text().isdigit():
            self.label_49.setText(u"Σφάλμα Συμπλήρωσης")
        else:
            self.label_49.setText(u"Καταχωρήθηκαν Επιτυχώς")

    def data(self):
        self.tables = []
        for table in self.document.tables:
            df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text:
                        df[i][j] = cell.text
            self.tables.append(pd.DataFrame(df))
            self.data = pd.concat(self.tables, axis=1)
            self.data.columns = ["A", "B", "C", "D", "E"]

    def second_data(self):
        self.second_tables = []
        for table in self.second_document.tables:
            df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text:
                        df[i][j] = cell.text
            self.second_tables.append(pd.DataFrame(df))
            self.second_data = pd.concat(self.second_tables, axis=1)
            self.second_data.columns = ["A", "B", "C", "D", "E", "F"]

        self.xamilo = self.second_document.paragraphs[11]
        self.meso_pros_xamilo = self.second_document.paragraphs[9]
        self.meso_pros_upsilo = self.second_document.paragraphs[7]
        self.upsilo = self.second_document.paragraphs[5]
        self.polu_upsilo = self.second_document.paragraphs[5]

    def holland_data(self):
        self.holland_tables = []
        for table in self.holland_document.tables:
            df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text:
                        df[i][j] = cell.text
            self.holland_tables.append(pd.DataFrame(df))

        self.drastiriotites_realstic = self.holland_tables[0]
        self.drastiriotites_investigative = self.holland_tables[1]
        self.drastiriotites_artistic = self.holland_tables[2]
        self.drastiriotites_social = self.holland_tables[3]
        self.drastiriotites_entreprising = self.holland_tables[4]
        self.drastiriotites_conventional = self.holland_tables[5]
        # data= pd.concat(tables, axis=1)
        self.ikanotites_realstic = self.holland_tables[6]
        self.ikanotites_investigative = self.holland_tables[7]
        self.ikanotites_artistic = self.holland_tables[8]
        self.ikanotites_social = self.holland_tables[9]
        self.ikanotites_entreprising = self.holland_tables[10]
        self.ikanotites_conventional = self.holland_tables[11]
        # rdvncxkzlw
        self.realstic = self.holland_tables[12]
        self.investigative = self.holland_tables[13]
        self.artistic = self.holland_tables[14]
        self.social = self.holland_tables[15]
        self.entreprising = self.holland_tables[16]
        self.conventional = self.holland_tables[17]
        #######
        self.self_assessment__1 = self.holland_tables[18]
        self.self_assessment__2 = self.holland_tables[19]

    def doc(self):
        self.doc_tables = []
        for table in self.doc1.tables:
            df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text:
                        df[i][j] = cell.text
            self.doc_tables.append(pd.DataFrame(df))
        self.que_1 = self.doc_tables[0]
        self.que_2 = self.doc_tables[1]
        self.que_drastiriotites = self.doc_tables[2]
        self.que_ikanotites = self.doc_tables[3]
        self.que_epaggelma = self.doc_tables[4]
        self.que_aksiologisi = self.doc_tables[5]

    def table1(self):
        self.model = TableModel(self.data)
        self.table_view.setModel(self.model)

        self.table_view.verticalHeader().setVisible(False)
        self.table_view.horizontalHeader().setVisible(False)
        self.table_view.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view.resizeColumnsToContents()
        # self.table_view.resizeRowsToContents()

        self.check1, self.check2, self.check3 = [], [], []
        # self.check3=[]
        for i, v in enumerate(self.data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem1 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem1)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check1.append(self.chkBoxItem1)
            self.table_view.setIndexWidget(self.table_view.model().index(i, 2), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem2 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem2)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check2.append(self.chkBoxItem2)
            self.table_view.setIndexWidget(self.table_view.model().index(i, 3), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem3 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem3)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check3.append(self.chkBoxItem3)
            self.table_view.setIndexWidget(self.table_view.model().index(i, 4), self.centerdCheckBoxWidget)

        self.checkBox_data = pd.DataFrame({"1": self.check1, "2": self.check2, "3": self.check3})

        a = []
        for i in range(len(self.data.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df = pd.DataFrame({"Label_1": a})
        self.df["Label_2"] = self.df["Label_1"]
        self.df["Label_3"] = self.df["Label_1"]
        self.df["Label_0"] = self.df["Label_1"]
        self.df["Label_4"] = self.df["Label_1"]
        self.df["Label_5"] = self.df["Label_1"]

    def checkboxChanged(self):
        self.labelResult1_1.setText("")
        for i, v in enumerate(self.checkBox_data["1"]):
            self.df["Label_1"][i] = QLabel()
            self.df["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_1.setText("{}, {}".format(self.labelResult1_1.text(),
                                                        self.df["Label_1"][i].text()))

        self.labelResult2_1.setText("")
        for i, v in enumerate(self.checkBox_data["2"]):
            self.df["Label_2"][i] = QLabel()
            self.df["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_1.setText("{}, {}".format(self.labelResult2_1.text(),
                                                        self.df["Label_2"][i].text()))

        self.labelResult3_1.setText("")
        for i, v in enumerate(self.checkBox_data["3"]):
            self.df["Label_3"][i] = QLabel()
            self.df["Label_3"][i].setText("True" if v.checkState() else "False")
            self.labelResult3_1.setText("{}, {}".format(self.labelResult3_1.text(),
                                                        self.df["Label_3"][i].text()))

        self.plot_list_1 = [int(self.labelResult1_1.text().count("True")),
                            int(self.labelResult2_1.text().count("True")),
                            int(self.labelResult3_1.text().count("True"))]

    def table2(self):
        self.data2 = self.data.copy()
        self.data2["B"] = ""
        self.data2["B"].iat[0] = "Ποιές σε εκφράζουν;"
        self.data2["C"].iat[0] = "Ποιές είναι καινούριες για σένα;"
        self.data2["D"].iat[0] = "Ποιές θα ήθελες να αποκτήσεις;"
        self.data2["E"].iat[0] = "Ποιές θα ήθελες να αλλάξεις;"

        self.model = TableModel(self.data2)
        self.table_view_2.setModel(self.model)

        self.table_view_2.verticalHeader().setVisible(False)
        self.table_view_2.horizontalHeader().setVisible(False)
        self.table_view_2.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_2.resizeColumnsToContents()

        self.check2_0, self.check2_1, self.check2_2, self.check2_3 = [], [], [], []
        # self.check3=[]
        for i, v in enumerate(self.data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem2_0 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem2_0)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check2_0.append(self.chkBoxItem2_0)
            self.table_view_2.setIndexWidget(self.table_view_2.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem2_1 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem2_1)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check2_1.append(self.chkBoxItem2_1)
            self.table_view_2.setIndexWidget(self.table_view_2.model().index(i, 2), self.centerdCheckBoxWidget)

        for i, v in enumerate(self.data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem2_2 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem2_2)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check2_2.append(self.chkBoxItem2_2)
            self.table_view_2.setIndexWidget(self.table_view_2.model().index(i, 3), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem2_3 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem2_3)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check2_3.append(self.chkBoxItem2_3)
            self.table_view_2.setIndexWidget(self.table_view_2.model().index(i, 4), self.centerdCheckBoxWidget)

        self.checkBox_data_2 = pd.DataFrame(
            {"0": self.check2_0, "1": self.check2_1, "2": self.check2_2, "3": self.check2_3})

    def checkboxChanged_2(self):
        self.labelResult0_2.setText("")
        for i, v in enumerate(self.checkBox_data_2["0"]):
            self.df["Label_0"][i] = QLabel()
            self.df["Label_0"][i].setText("True" if v.checkState() else "False")
            self.labelResult0_2.setText("{}, {}".format(self.labelResult0_2.text(),
                                                        self.df["Label_0"][i].text()))
        self.labelResult1_2.setText("")
        for i, v in enumerate(self.checkBox_data_2["1"]):
            self.df["Label_1"][i] = QLabel()
            self.df["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_2.setText("{}, {}".format(self.labelResult1_2.text(),
                                                        self.df["Label_1"][i].text()))
        self.labelResult2_2.setText("")
        for i, v in enumerate(self.checkBox_data_2["2"]):
            self.df["Label_2"][i] = QLabel()
            self.df["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_2.setText("{}, {}".format(self.labelResult2_2.text(),
                                                        self.df["Label_2"][i].text()))
        self.labelResult3_2.setText("")
        for i, v in enumerate(self.checkBox_data_2["3"]):
            self.df["Label_3"][i] = QLabel()
            self.df["Label_3"][i].setText("True" if v.checkState() else "False")
            self.labelResult3_2.setText("{}, {}".format(self.labelResult3_2.text(),
                                                        self.df["Label_3"][i].text()))
        self.plot_list_2 = [int(self.labelResult0_2.text().count("True")),
                            int(self.labelResult1_2.text().count("True")),
                            int(self.labelResult2_2.text().count("True")),
                            int(self.labelResult3_2.text().count("True"))]

    def table1_results(self):
        a = self.labelResult0_2.text()
        b = self.labelResult1_2.text()
        c = self.labelResult2_2.text()
        d = self.labelResult3_2.text()

        y = self.data2["A"].values.tolist()
        y.pop(0)

        a0 = a.split(",")
        a0.pop(0)
        x1 = [1 if i == " True" else 0 for i in a0]
        self.x1 = ", ".join(self.cross(x1, y))
        self.x11 = u"Οι αξίες που εκφράζουν τον ενδιαφερόμενο είναι οι εξής: "
        self.label_51.setText(self.x11 + self.x1)

        b0 = b.split(",")
        b0.pop(0)
        x2 = [1 if i == " True" else 0 for i in b0]
        self.x2 = ", ".join(self.cross(x2, y))
        if (self.x2 == ""):
            self.x22_0 = u"O ενδιαφερόμενος γνωρίζει όλες τις αξίες"
            self.label_52.setText(self.x22_0)
        else:
            self.x22 = u"Οι αξίες που είναι και καινούριες για τον ενδιαφερόμενο είναι οι εξής: "
            self.label_52.setText(self.x22 + self.x2)

        c0 = c.split(",")
        c0.pop(0)
        x3 = [1 if i == " True" else 0 for i in c0]
        self.x3 = ", ".join(self.cross(x3, y))
        self.x33 = u"Οι αξίες που θα ήθελε να αποκτήσει ο ενδιαφερόμενος είναι οι εξής: "
        self.label_53.setText(self.x33 + self.x3)

        d0 = d.split(",")
        d0.pop(0)
        x4 = [1 if i == " True" else 0 for i in d0]
        self.x4 = ", ".join(self.cross(x4, y))
        self.x44 = u"Οι αξίες που θα ήθελε να αλλάξει ο ενδιαφερόμενος είναι οι εξής: "
        self.label_54.setText(self.x44 + self.x4)

    def update_graph_1(self):

        labels = 'Ποτέ', 'Συχνά', 'Πάντα'
        explode = (0.1, 0.1, 0.1)  # only "explode" the 2nd slice (i.e. 'Hogs')
        # Creating color scale
        cmap = plt.get_cmap("tab20c")
        colors = cmap(np.array([12, 13, 15]))
        fig1, ax1 = plt.subplots()
        ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a cir
        ax1.set_title(u"Ερωτηματολόγιο Αξιών")
        self.MplWidget.canvas.axes.clear()
        self.MplWidget.canvas.axes.pie(self.plot_list_1, explode=explode, labels=labels, autopct='%1.1f%%',
                                       shadow=True, startangle=90, colors=colors)

        self.MplWidget.canvas.draw()

    def update_graph_1_2(self):

        labels = ['Ποτέ', 'Συχνά', 'Πάντα']

        fig1, ax1 = plt.subplots()

        self.MplWidget_6.canvas.axes.clear()
        self.MplWidget_6.canvas.axes.bar(labels, self.plot_list_1, color='mediumpurple', width=0.2)
        ax1.axis(ymin=0, ymax=35)
        ax1.set_ylabel(u"Σύνολο Ερωτήσεων")
        self.MplWidget_6.canvas.draw()

    def update_graph_2(self):

        labels = "Με Εκφράζουν", 'Νέες', 'Να Αποκτηθούν', 'Να Aλλάξουν'

        explode = (0.1, 0.1, 0.1, 0.1)
        # Creating color scale
        cmap = plt.get_cmap("tab20c")
        colors = cmap(np.array([12, 13, 15]))

        fig1, ax1 = plt.subplots()
        # ax1.pie(self.plot_list_2, explode=explode, labels=labels, autopct='%1.1f%%',
        #         shadow=True, startangle=90)
        ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a cir

        self.MplWidget_4.canvas.axes.clear()
        self.MplWidget_4.canvas.axes.pie(self.plot_list_2, explode=explode, labels=labels, autopct='%1.1f%%',
                                         shadow=True, startangle=90, colors=colors)

        self.MplWidget_4.canvas.draw()

    def table3(self):
        self.model = TableModel(self.second_data)
        self.table_view_3.setModel(self.model)

        self.table_view_3.verticalHeader().setVisible(False)
        self.table_view_3.horizontalHeader().setVisible(False)
        self.table_view_3.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_3.resizeColumnsToContents()
        # self.table_view.resizeRowsToContents()

        self.check3_1, self.check3_2, self.check3_3, self.check3_4, self.check3_5 = [], [], [], [], []
        # self.check3=[]
        for i, v in enumerate(self.second_data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem3_1 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem3_1)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check3_1.append(self.chkBoxItem3_1)
            self.table_view_3.setIndexWidget(self.table_view_3.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.second_data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem3_2 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem3_2)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check3_2.append(self.chkBoxItem3_2)
            self.table_view_3.setIndexWidget(self.table_view_3.model().index(i, 2), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.second_data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem3_3 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem3_3)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check3_3.append(self.chkBoxItem3_3)
            self.table_view_3.setIndexWidget(self.table_view_3.model().index(i, 3), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.second_data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem3_4 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem3_4)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check3_4.append(self.chkBoxItem3_4)
            self.table_view_3.setIndexWidget(self.table_view_3.model().index(i, 4), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.second_data.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem3_5 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem3_5)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check3_5.append(self.chkBoxItem3_5)
            self.table_view_3.setIndexWidget(self.table_view_3.model().index(i, 5), self.centerdCheckBoxWidget)

        self.checkBox_data_3 = pd.DataFrame(
            {"1": self.check3_1, "2": self.check3_2, "3": self.check3_3, "4": self.check3_4, "5": self.check3_5})
        a = []
        for i in range(len(self.second_data.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df2 = pd.DataFrame({"Label_1": a})
        self.df2["Label_2"] = self.df2["Label_1"]
        self.df2["Label_3"] = self.df2["Label_1"]
        self.df2["Label_0"] = self.df2["Label_1"]
        self.df2["Label_4"] = self.df2["Label_1"]
        self.df2["Label_5"] = self.df2["Label_1"]

    def checkboxChanged_3(self):
        self.labelResult1_3.setText("")
        for i, v in enumerate(self.checkBox_data_3["1"]):
            self.df2["Label_1"][i] = QLabel()
            self.df2["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_3.setText("{}, {}".format(self.labelResult1_3.text(),
                                                        self.df2["Label_1"][i].text()))

        self.labelResult2_3.setText("")
        for i, v in enumerate(self.checkBox_data_3["2"]):
            self.df2["Label_2"][i] = QLabel()
            self.df2["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_3.setText("{}, {}".format(self.labelResult2_3.text(),
                                                        self.df2["Label_2"][i].text()))

        self.labelResult3_3.setText("")
        for i, v in enumerate(self.checkBox_data_3["3"]):
            self.df2["Label_3"][i] = QLabel()
            self.df2["Label_3"][i].setText("True" if v.checkState() else "False")
            self.labelResult3_3.setText("{}, {}".format(self.labelResult3_3.text(),
                                                        self.df2["Label_3"][i].text()))
        self.labelResult4_3.setText("")
        for i, v in enumerate(self.checkBox_data_3["4"]):
            self.df2["Label_4"][i] = QLabel()
            self.df2["Label_4"][i].setText("True" if v.checkState() else "False")
            self.labelResult4_3.setText("{}, {}".format(self.labelResult4_3.text(),
                                                        self.df2["Label_4"][i].text()))

        self.labelResult5_3.setText("")
        for i, v in enumerate(self.checkBox_data_3["5"]):
            self.df2["Label_5"][i] = QLabel()
            self.df2["Label_5"][i].setText("True" if v.checkState() else "False")
            self.labelResult5_3.setText("{}, {}".format(self.labelResult5_3.text(),
                                                        self.df2["Label_5"][i].text()))
        self.plot_list_3 = [int(self.labelResult1_3.text().count("True")),
                            int(self.labelResult2_3.text().count("True")),
                            int(self.labelResult3_3.text().count("True")),
                            int(self.labelResult4_3.text().count("True")),
                            int(self.labelResult5_3.text().count("True"))]

        # print(self.sum)

    def table1_results_2(self):
        # pass
        self.list_3 = [int(self.labelResult1_3.text().count("True")) * 0,
                       int(self.labelResult2_3.text().count("True")) * 1,
                       int(self.labelResult3_3.text().count("True")) * 2,
                       int(self.labelResult4_3.text().count("True")) * 3,
                       int(self.labelResult5_3.text().count("True")) * 4]
        # print(self.list_3)
        self.sum = sum(self.list_3)

        self.y1 = u"Ο Ενδιαφερόμενος εμφάνησε: "
        if (int(self.lineEdit_2.text()) <= 16):
            if (self.sum <= 8):
                self.y11 = u"χαμηλό"
            elif (self.sum >= 9 and self.sum <= 17):
                self.y11 = u"μέσο προς χαμηλό"
            elif (self.sum >= 18 and self.sum <= 33):
                self.y11 = u"μέσο προς υψηλό"
            elif (self.sum >= 34 and self.sum <= 54):
                self.y11 = u"υψηλό"
            else:
                self.y11 = u"πολύ υψηλό"
        elif (int(self.lineEdit_2.text()) >= 17 and int(self.lineEdit_2.text()) <= 21):
            if self.sum <= 20:
                self.y11 = u"χαμηλό"
            elif (self.sum >= 21 and self.sum <= 36):
                self.y11 = u"μέσο προς χαμηλό"
            elif (self.sum >= 37 and self.sum <= 44):
                self.y11 = u"μέσο προς υψηλό"
            elif (self.sum >= 45 and self.sum <= 69):
                self.y11 = u"υψηλό"
            else:
                self.y11 = u"πολύ υψηλό"
        elif (int(self.lineEdit_2.text()) >= 22 and int(self.lineEdit_2.text()) <= 30):
            if self.sum <= 12:
                self.y11 = u"χαμηλό"
            elif (self.sum >= 13 and self.sum <= 25):
                self.y11 = u"μέσο προς χαμηλό"
            elif (self.sum >= 26 and self.sum <= 40):
                self.y11 = u"μέσο προς υψηλό"
            elif (self.sum >= 41 and self.sum <= 59):
                self.y11 = u"υψηλό"
            else:
                self.y11 = u"πολύ υψηλό"
        else:
            if self.sum <= 15:
                self.y11 = u"χαμηλό"
            elif (self.sum >= 16 and self.sum <= 29):
                self.y11 = u"μέσο προς χαμηλό"
            elif (self.sum >= 30 and self.sum <= 46):
                self.y11 = u"μέσο προς υψηλό"
            elif (self.sum >= 47 and self.sum <= 66):
                self.y11 = u"υψηλό"
            else:
                self.y11 = u"πολύ υψηλό"

        self.y111 = u" επίπεδο αυτοπεποίθησης."

        self.label_31.setText(self.y1 + "\u0332".join(self.y11) + self.y111)

        if (self.y11 == u"χαμηλό"):
            self.textEdit.setText(self.xamilo.text)
        elif (self.y11 == u"μέσο προς χαμηλό"):
            self.textEdit.setText(self.meso_pros_xamilo.text)
        elif (self.y11 == u"μέσο προς υψηλό"):
            self.textEdit.setText(self.meso_pros_upsilo.text)
        elif (self.y11 == u"υψηλό"):
            self.textEdit.setText(self.upsilo.text)
        else:
            self.textEdit.setText(self.polu_upsilo.text)

    def update_graph_3(self):

        labels = 'Ποτέ', "Σπάνια", "Ενίοτε", 'Συχνά', 'Πάντα'

        explode = (0.1, 0.1, 0.1, 0.1, 0.1)
        # Creating color scale
        cmap = plt.get_cmap("tab20c")
        colors = cmap(np.array([12, 13, 15]))

        fig1, ax1 = plt.subplots()
        ax1.pie(self.plot_list_3, explode=explode, labels=labels, autopct='%1.1f%%',
                shadow=True, startangle=90)
        ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a cir

        self.MplWidget_2.canvas.axes.clear()
        self.MplWidget_2.canvas.axes.pie(self.plot_list_3, explode=explode, labels=labels, autopct='%1.1f%%',
                                         shadow=True, startangle=90, colors=colors)

        self.MplWidget_2.canvas.draw()

    def update_graph_3_1(self):

        labels = 'Ποτέ', "Σπάνια", "Ενίοτε", 'Συχνά', 'Πάντα'

        fig1, ax1 = plt.subplots()

        self.MplWidget_7.canvas.axes.clear()
        self.MplWidget_7.canvas.axes.bar(labels, self.plot_list_3, color='mediumpurple', width=0.2)
        ax1.axis(ymin=0, ymax=35)
        ax1.set_ylabel(u"Σύνολο Ερωτήσεων")
        self.MplWidget_7.canvas.draw()

    def table4_1(self):  # dokimio -> drastiriotites
        self.model = TableModel(self.drastiriotites_realstic)
        self.table_view_7.setModel(self.model)

        self.table_view_7.verticalHeader().setVisible(False)
        self.table_view_7.horizontalHeader().setVisible(False)
        self.table_view_7.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_7.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_1, self.check4_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.drastiriotites_realstic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_1 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_1)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_1.append(self.chkBoxItem4_1)
            self.table_view_7.setIndexWidget(self.table_view_7.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.drastiriotites_realstic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_2 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_2)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_2.append(self.chkBoxItem4_2)
            self.table_view_7.setIndexWidget(self.table_view_7.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_1 = pd.DataFrame({"1": self.check4_1, "2": self.check4_2})
        a = []
        for i in range(len(self.drastiriotites_realstic.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_1 = pd.DataFrame({"Label_1": a})
        self.df4_1["Label_2"] = self.df4_1["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_1(self):
        self.labelResult1_4_1.setText("")
        for i, v in enumerate(self.checkBox_data_4_1["1"]):
            self.df4_1["Label_1"][i] = QLabel()
            self.df4_1["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_1.setText("{}, {}".format(self.labelResult1_4_1.text(),
                                                          self.df4_1["Label_1"][i].text()))
        self.labelResult2_4_1.setText("")
        for i, v in enumerate(self.checkBox_data_4_1["2"]):
            self.df4_1["Label_2"][i] = QLabel()
            self.df4_1["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_1.setText("{}, {}".format(self.labelResult2_4_1.text(),
                                                          self.df4_1["Label_2"][i].text()))

        self.plot_list_4_1 = [int(self.labelResult1_4_1.text().count("True")),
                              int(self.labelResult2_4_1.text().count("True"))]

    def table4_2(self):  # dokimio -> drastiriotites
        self.model = TableModel(self.drastiriotites_investigative)
        self.table_view_8.setModel(self.model)

        self.table_view_8.verticalHeader().setVisible(False)
        self.table_view_8.horizontalHeader().setVisible(False)
        self.table_view_8.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_8.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_2_1, self.check4_2_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.drastiriotites_investigative.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_3 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_3)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_2_1.append(self.chkBoxItem4_3)
            self.table_view_8.setIndexWidget(self.table_view_8.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.drastiriotites_investigative.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_4 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_4)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_2_2.append(self.chkBoxItem4_4)
            self.table_view_8.setIndexWidget(self.table_view_8.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_2 = pd.DataFrame({"1": self.check4_2_1, "2": self.check4_2_2})
        a = []
        for i in range(len(self.drastiriotites_realstic.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_2 = pd.DataFrame({"Label_1": a})
        self.df4_2["Label_2"] = self.df4_2["Label_1"]

    def checkboxChanged_4_2(self):
        self.labelResult1_4_2.setText("")
        for i, v in enumerate(self.checkBox_data_4_2["1"]):
            self.df4_2["Label_1"][i] = QLabel()
            self.df4_2["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_2.setText("{}, {}".format(self.labelResult1_4_2.text(),
                                                          self.df4_2["Label_1"][i].text()))
        self.labelResult2_4_2.setText("")
        for i, v in enumerate(self.checkBox_data_4_2["2"]):
            self.df4_2["Label_2"][i] = QLabel()
            self.df4_2["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_2.setText("{}, {}".format(self.labelResult2_4_2.text(),
                                                          self.df4_2["Label_2"][i].text()))

        self.plot_list_4_2 = [int(self.labelResult1_4_2.text().count("True")),
                              int(self.labelResult2_4_2.text().count("True"))]

    def table4_3(self):  # dokimio -> drastiriotites
        self.model = TableModel(self.drastiriotites_artistic)
        self.table_view_9.setModel(self.model)

        self.table_view_9.verticalHeader().setVisible(False)
        self.table_view_9.horizontalHeader().setVisible(False)
        self.table_view_9.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_9.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_3_1, self.check4_3_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.drastiriotites_artistic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_5 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_5)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_3_1.append(self.chkBoxItem4_5)
            self.table_view_9.setIndexWidget(self.table_view_9.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.drastiriotites_artistic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_6 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_6)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_3_2.append(self.chkBoxItem4_6)
            self.table_view_9.setIndexWidget(self.table_view_9.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_3 = pd.DataFrame({"1": self.check4_3_1, "2": self.check4_3_2})
        a = []
        for i in range(len(self.drastiriotites_artistic.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_3 = pd.DataFrame({"Label_1": a})
        self.df4_3["Label_2"] = self.df4_3["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_3(self):
        self.labelResult1_4_3.setText("")
        for i, v in enumerate(self.checkBox_data_4_3["1"]):
            self.df4_3["Label_1"][i] = QLabel()
            self.df4_3["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_3.setText("{}, {}".format(self.labelResult1_4_3.text(),
                                                          self.df4_3["Label_1"][i].text()))
        self.labelResult2_4_3.setText("")
        for i, v in enumerate(self.checkBox_data_4_3["2"]):
            self.df4_3["Label_2"][i] = QLabel()
            self.df4_3["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_3.setText("{}, {}".format(self.labelResult2_4_3.text(),
                                                          self.df4_3["Label_2"][i].text()))

        self.plot_list_4_3 = [int(self.labelResult1_4_3.text().count("True")),
                              int(self.labelResult2_4_3.text().count("True"))]

        # print(self.checkBox_data_3)

    def table4_4(self):  # dokimio -> drastiriotites
        self.model = TableModel(self.drastiriotites_social)
        self.table_view_10.setModel(self.model)

        self.table_view_10.verticalHeader().setVisible(False)
        self.table_view_10.horizontalHeader().setVisible(False)
        self.table_view_10.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_10.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_4_1, self.check4_4_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.drastiriotites_social.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_7 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_7)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_4_1.append(self.chkBoxItem4_7)
            self.table_view_10.setIndexWidget(self.table_view_10.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.drastiriotites_social.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_8 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_8)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_4_2.append(self.chkBoxItem4_8)
            self.table_view_10.setIndexWidget(self.table_view_10.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_4 = pd.DataFrame({"1": self.check4_4_1, "2": self.check4_4_2})
        a = []
        for i in range(len(self.drastiriotites_artistic.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_4 = pd.DataFrame({"Label_1": a})
        self.df4_4["Label_2"] = self.df4_4["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_4(self):
        self.labelResult1_4_4.setText("")
        for i, v in enumerate(self.checkBox_data_4_4["1"]):
            self.df4_4["Label_1"][i] = QLabel()
            self.df4_4["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_4.setText("{}, {}".format(self.labelResult1_4_4.text(),
                                                          self.df4_4["Label_1"][i].text()))
        self.labelResult2_4_4.setText("")
        for i, v in enumerate(self.checkBox_data_4_4["2"]):
            self.df4_4["Label_2"][i] = QLabel()
            self.df4_4["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_4.setText("{}, {}".format(self.labelResult2_4_4.text(),
                                                          self.df4_4["Label_2"][i].text()))

        self.plot_list_4_4 = [int(self.labelResult1_4_4.text().count("True")),
                              int(self.labelResult2_4_4.text().count("True"))]

    def table4_5(self):  # dokimio -> drastiriotites
        self.model = TableModel(self.drastiriotites_entreprising)
        self.table_view_11.setModel(self.model)

        self.table_view_11.verticalHeader().setVisible(False)
        self.table_view_11.horizontalHeader().setVisible(False)
        self.table_view_11.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_11.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_5_1, self.check4_5_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.drastiriotites_entreprising.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_9 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_9)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_5_1.append(self.chkBoxItem4_9)
            self.table_view_11.setIndexWidget(self.table_view_11.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.drastiriotites_entreprising.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_10 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_10)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_5_2.append(self.chkBoxItem4_10)
            self.table_view_11.setIndexWidget(self.table_view_11.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_5 = pd.DataFrame({"1": self.check4_5_1, "2": self.check4_5_2})
        a = []
        for i in range(len(self.drastiriotites_entreprising.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_5 = pd.DataFrame({"Label_1": a})
        self.df4_5["Label_2"] = self.df4_5["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_5(self):
        self.labelResult1_4_5.setText("")
        for i, v in enumerate(self.checkBox_data_4_5["1"]):
            self.df4_5["Label_1"][i] = QLabel()
            self.df4_5["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_5.setText("{}, {}".format(self.labelResult1_4_5.text(),
                                                          self.df4_5["Label_1"][i].text()))
        self.labelResult2_4_5.setText("")
        for i, v in enumerate(self.checkBox_data_4_5["2"]):
            self.df4_5["Label_2"][i] = QLabel()
            self.df4_5["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_5.setText("{}, {}".format(self.labelResult2_4_5.text(),
                                                          self.df4_5["Label_2"][i].text()))

        self.plot_list_4_5 = [int(self.labelResult1_4_5.text().count("True")),
                              int(self.labelResult2_4_5.text().count("True"))]

    def table4_6(self):  # dokimio -> drastiriotites
        self.model = TableModel(self.drastiriotites_conventional)
        self.table_view_12.setModel(self.model)

        self.table_view_12.verticalHeader().setVisible(False)
        self.table_view_12.horizontalHeader().setVisible(False)
        self.table_view_12.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_12.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_6_1, self.check4_6_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.drastiriotites_conventional.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_11 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_11)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_6_1.append(self.chkBoxItem4_11)
            self.table_view_12.setIndexWidget(self.table_view_12.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.drastiriotites_conventional.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_12 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_12)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_6_2.append(self.chkBoxItem4_12)
            self.table_view_12.setIndexWidget(self.table_view_12.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_6 = pd.DataFrame({"1": self.check4_6_1, "2": self.check4_6_2})
        a = []
        for i in range(len(self.drastiriotites_conventional.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_6 = pd.DataFrame({"Label_1": a})
        self.df4_6["Label_2"] = self.df4_6["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_6(self):
        self.labelResult1_4_6.setText("")
        for i, v in enumerate(self.checkBox_data_4_6["1"]):
            self.df4_6["Label_1"][i] = QLabel()
            self.df4_6["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_6.setText("{}, {}".format(self.labelResult1_4_6.text(),
                                                          self.df4_6["Label_1"][i].text()))
        self.labelResult2_4_6.setText("")
        for i, v in enumerate(self.checkBox_data_4_6["2"]):
            self.df4_6["Label_2"][i] = QLabel()
            self.df4_6["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_6.setText("{}, {}".format(self.labelResult2_4_6.text(),
                                                          self.df4_6["Label_2"][i].text()))

        self.plot_list_4_6 = [int(self.labelResult1_4_6.text().count("True")),
                              int(self.labelResult2_4_6.text().count("True"))]

    def table4_7(self):  # neo tab
        self.model = TableModel(self.ikanotites_realstic)
        self.table_view_13.setModel(self.model)

        self.table_view_13.verticalHeader().setVisible(False)
        self.table_view_13.horizontalHeader().setVisible(False)
        self.table_view_13.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_13.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_7_1, self.check4_7_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.ikanotites_realstic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_13 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_13)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_7_1.append(self.chkBoxItem4_13)
            self.table_view_13.setIndexWidget(self.table_view_13.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.ikanotites_realstic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_14 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_14)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_7_2.append(self.chkBoxItem4_14)
            self.table_view_13.setIndexWidget(self.table_view_13.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_7 = pd.DataFrame({"1": self.check4_7_1, "2": self.check4_7_2})
        a = []
        for i in range(len(self.ikanotites_realstic.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_7 = pd.DataFrame({"Label_1": a})
        self.df4_7["Label_2"] = self.df4_7["Label_1"]

    def checkboxChanged_4_7(self):
        self.labelResult1_4_7.setText("")
        for i, v in enumerate(self.checkBox_data_4_7["1"]):
            self.df4_7["Label_1"][i] = QLabel()
            self.df4_7["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_7.setText("{}, {}".format(self.labelResult1_4_7.text(),
                                                          self.df4_7["Label_1"][i].text()))
        self.labelResult2_4_7.setText("")
        for i, v in enumerate(self.checkBox_data_4_7["2"]):
            self.df4_7["Label_2"][i] = QLabel()
            self.df4_7["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_7.setText("{}, {}".format(self.labelResult2_4_7.text(),
                                                          self.df4_7["Label_2"][i].text()))

        self.plot_list_4_7 = [int(self.labelResult1_4_7.text().count("True")),
                              int(self.labelResult2_4_7.text().count("True"))]

    def table4_8(self):  # neo tab
        self.model = TableModel(self.ikanotites_investigative)
        self.table_view_14.setModel(self.model)

        self.table_view_14.verticalHeader().setVisible(False)
        self.table_view_14.horizontalHeader().setVisible(False)
        self.table_view_14.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_14.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_8_1, self.check4_8_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.ikanotites_investigative.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_15 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_15)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_8_1.append(self.chkBoxItem4_15)
            self.table_view_14.setIndexWidget(self.table_view_14.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.ikanotites_investigative.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_16 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_16)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_8_2.append(self.chkBoxItem4_16)
            self.table_view_14.setIndexWidget(self.table_view_14.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_8 = pd.DataFrame({"1": self.check4_8_1, "2": self.check4_8_2})
        a = []
        for i in range(len(self.ikanotites_investigative.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_8 = pd.DataFrame({"Label_1": a})
        self.df4_8["Label_2"] = self.df4_8["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_8(self):
        self.labelResult1_4_8.setText("")
        for i, v in enumerate(self.checkBox_data_4_8["1"]):
            self.df4_8["Label_1"][i] = QLabel()
            self.df4_8["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_8.setText("{}, {}".format(self.labelResult1_4_8.text(),
                                                          self.df4_8["Label_1"][i].text()))
        self.labelResult2_4_8.setText("")
        for i, v in enumerate(self.checkBox_data_4_8["2"]):
            self.df4_8["Label_2"][i] = QLabel()
            self.df4_8["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_8.setText("{}, {}".format(self.labelResult2_4_8.text(),
                                                          self.df4_8["Label_2"][i].text()))

        self.plot_list_4_8 = [int(self.labelResult1_4_8.text().count("True")),
                              int(self.labelResult2_4_8.text().count("True"))]

    def table4_9(self):  # neo tab
        self.model = TableModel(self.ikanotites_artistic)
        self.table_view_15.setModel(self.model)

        self.table_view_15.verticalHeader().setVisible(False)
        self.table_view_15.horizontalHeader().setVisible(False)
        self.table_view_15.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_15.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_9_1, self.check4_9_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.ikanotites_artistic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_17 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_17)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_9_1.append(self.chkBoxItem4_17)
            self.table_view_15.setIndexWidget(self.table_view_15.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.ikanotites_artistic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_18 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_18)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_9_2.append(self.chkBoxItem4_18)
            self.table_view_15.setIndexWidget(self.table_view_15.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_9 = pd.DataFrame({"1": self.check4_9_1, "2": self.check4_9_2})
        a = []
        for i in range(len(self.ikanotites_artistic.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_9 = pd.DataFrame({"Label_1": a})
        self.df4_9["Label_2"] = self.df4_9["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_9(self):
        self.labelResult1_4_9.setText("")
        for i, v in enumerate(self.checkBox_data_4_9["1"]):
            self.df4_9["Label_1"][i] = QLabel()
            self.df4_9["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_9.setText("{}, {}".format(self.labelResult1_4_9.text(),
                                                          self.df4_9["Label_1"][i].text()))
        self.labelResult2_4_9.setText("")
        for i, v in enumerate(self.checkBox_data_4_9["2"]):
            self.df4_9["Label_2"][i] = QLabel()
            self.df4_9["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_9.setText("{}, {}".format(self.labelResult2_4_9.text(),
                                                          self.df4_9["Label_2"][i].text()))

        self.plot_list_4_9 = [int(self.labelResult1_4_9.text().count("True")),
                              int(self.labelResult2_4_9.text().count("True"))]

    def table4_10(self):  # neo tab
        self.model = TableModel(self.ikanotites_social)
        self.table_view_16.setModel(self.model)

        self.table_view_16.verticalHeader().setVisible(False)
        self.table_view_16.horizontalHeader().setVisible(False)
        self.table_view_16.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_16.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_10_1, self.check4_10_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.ikanotites_social.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_19 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_19)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_10_1.append(self.chkBoxItem4_19)
            self.table_view_16.setIndexWidget(self.table_view_16.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.ikanotites_social.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_20 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_20)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_10_2.append(self.chkBoxItem4_20)
            self.table_view_16.setIndexWidget(self.table_view_16.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_10 = pd.DataFrame({"1": self.check4_10_1, "2": self.check4_10_2})
        a = []
        for i in range(len(self.ikanotites_social.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_10 = pd.DataFrame({"Label_1": a})
        self.df4_10["Label_2"] = self.df4_10["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_10(self):
        self.labelResult1_4_10.setText("")
        for i, v in enumerate(self.checkBox_data_4_10["1"]):
            self.df4_10["Label_1"][i] = QLabel()
            self.df4_10["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_10.setText("{}, {}".format(self.labelResult1_4_10.text(),
                                                           self.df4_10["Label_1"][i].text()))
        self.labelResult2_4_10.setText("")
        for i, v in enumerate(self.checkBox_data_4_10["2"]):
            self.df4_10["Label_2"][i] = QLabel()
            self.df4_10["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_10.setText("{}, {}".format(self.labelResult2_4_10.text(),
                                                           self.df4_10["Label_2"][i].text()))

        self.plot_list_4_10 = [int(self.labelResult1_4_10.text().count("True")),
                               int(self.labelResult2_4_10.text().count("True"))]

    def table4_11(self):  # neo tab
        self.model = TableModel(self.ikanotites_entreprising)
        self.table_view_17.setModel(self.model)

        self.table_view_17.verticalHeader().setVisible(False)
        self.table_view_17.horizontalHeader().setVisible(False)
        self.table_view_17.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_17.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_11_1, self.check4_11_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.ikanotites_entreprising.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_21 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_21)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_11_1.append(self.chkBoxItem4_21)
            self.table_view_17.setIndexWidget(self.table_view_17.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.ikanotites_entreprising.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_22 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_22)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_11_2.append(self.chkBoxItem4_22)
            self.table_view_17.setIndexWidget(self.table_view_17.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_11 = pd.DataFrame({"1": self.check4_11_1, "2": self.check4_11_2})
        a = []
        for i in range(len(self.ikanotites_entreprising.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_11 = pd.DataFrame({"Label_1": a})
        self.df4_11["Label_2"] = self.df4_11["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_11(self):
        self.labelResult1_4_11.setText("")
        for i, v in enumerate(self.checkBox_data_4_11["1"]):
            self.df4_11["Label_1"][i] = QLabel()
            self.df4_11["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_11.setText("{}, {}".format(self.labelResult1_4_11.text(),
                                                           self.df4_11["Label_1"][i].text()))
        self.labelResult2_4_11.setText("")
        for i, v in enumerate(self.checkBox_data_4_11["2"]):
            self.df4_11["Label_2"][i] = QLabel()
            self.df4_11["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_11.setText("{}, {}".format(self.labelResult2_4_11.text(),
                                                           self.df4_11["Label_2"][i].text()))

        self.plot_list_4_11 = [int(self.labelResult1_4_11.text().count("True")),
                               int(self.labelResult2_4_11.text().count("True"))]

    def table4_12(self):  # neo tab
        self.model = TableModel(self.ikanotites_conventional)
        self.table_view_18.setModel(self.model)

        self.table_view_18.verticalHeader().setVisible(False)
        self.table_view_18.horizontalHeader().setVisible(False)
        self.table_view_18.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_18.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_12_1, self.check4_12_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.ikanotites_conventional.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_23 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_23)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_12_1.append(self.chkBoxItem4_23)
            self.table_view_18.setIndexWidget(self.table_view_18.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.ikanotites_conventional.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_24 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_24)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_12_2.append(self.chkBoxItem4_24)
            self.table_view_18.setIndexWidget(self.table_view_18.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_12 = pd.DataFrame({"1": self.check4_12_1, "2": self.check4_12_2})
        a = []
        for i in range(len(self.ikanotites_conventional.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_12 = pd.DataFrame({"Label_1": a})
        self.df4_12["Label_2"] = self.df4_12["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_12(self):
        self.labelResult1_4_12.setText("")
        for i, v in enumerate(self.checkBox_data_4_12["1"]):
            self.df4_12["Label_1"][i] = QLabel()
            self.df4_12["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_12.setText("{}, {}".format(self.labelResult1_4_12.text(),
                                                           self.df4_12["Label_1"][i].text()))
        self.labelResult2_4_12.setText("")
        for i, v in enumerate(self.checkBox_data_4_12["2"]):
            self.df4_12["Label_2"][i] = QLabel()
            self.df4_12["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_12.setText("{}, {}".format(self.labelResult2_4_12.text(),
                                                           self.df4_12["Label_2"][i].text()))

        self.plot_list_4_12 = [int(self.labelResult1_4_12.text().count("True")),
                               int(self.labelResult2_4_12.text().count("True"))]

    def table4_13(self):  # neo tab 2
        self.model = TableModel(self.realstic)
        self.table_view_19.setModel(self.model)

        self.table_view_19.verticalHeader().setVisible(False)
        self.table_view_19.horizontalHeader().setVisible(False)
        self.table_view_19.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_19.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_13_1, self.check4_13_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.realstic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_25 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_25)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_13_1.append(self.chkBoxItem4_25)
            self.table_view_19.setIndexWidget(self.table_view_19.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.realstic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_26 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_26)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_13_2.append(self.chkBoxItem4_26)
            self.table_view_19.setIndexWidget(self.table_view_19.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_13 = pd.DataFrame({"1": self.check4_13_1, "2": self.check4_13_2})
        a = []
        for i in range(len(self.realstic.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_13 = pd.DataFrame({"Label_1": a})
        self.df4_13["Label_2"] = self.df4_13["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_13(self):
        self.labelResult1_4_13.setText("")
        for i, v in enumerate(self.checkBox_data_4_13["1"]):
            self.df4_13["Label_1"][i] = QLabel()
            self.df4_13["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_13.setText("{}, {}".format(self.labelResult1_4_13.text(),
                                                           self.df4_13["Label_1"][i].text()))
        self.labelResult2_4_13.setText("")
        for i, v in enumerate(self.checkBox_data_4_13["2"]):
            self.df4_13["Label_2"][i] = QLabel()
            self.df4_13["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_13.setText("{}, {}".format(self.labelResult2_4_13.text(),
                                                           self.df4_13["Label_2"][i].text()))

        self.plot_list_4_13 = [int(self.labelResult1_4_13.text().count("True")),
                               int(self.labelResult2_4_13.text().count("True"))]

    def table4_14(self):  # neo tab 2
        self.model = TableModel(self.investigative)
        self.table_view_20.setModel(self.model)

        self.table_view_20.verticalHeader().setVisible(False)
        self.table_view_20.horizontalHeader().setVisible(False)
        self.table_view_20.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_20.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_14_1, self.check4_14_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.investigative.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_27 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_27)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_14_1.append(self.chkBoxItem4_27)
            self.table_view_20.setIndexWidget(self.table_view_20.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.investigative.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_28 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_28)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_14_2.append(self.chkBoxItem4_28)
            self.table_view_20.setIndexWidget(self.table_view_20.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_14 = pd.DataFrame({"1": self.check4_14_1, "2": self.check4_14_2})
        a = []
        for i in range(len(self.investigative.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_14 = pd.DataFrame({"Label_1": a})
        self.df4_14["Label_2"] = self.df4_14["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_14(self):
        self.labelResult1_4_14.setText("")
        for i, v in enumerate(self.checkBox_data_4_14["1"]):
            self.df4_14["Label_1"][i] = QLabel()
            self.df4_14["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_14.setText("{}, {}".format(self.labelResult1_4_14.text(),
                                                           self.df4_14["Label_1"][i].text()))
        self.labelResult2_4_14.setText("")
        for i, v in enumerate(self.checkBox_data_4_14["2"]):
            self.df4_14["Label_2"][i] = QLabel()
            self.df4_14["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_14.setText("{}, {}".format(self.labelResult2_4_14.text(),
                                                           self.df4_14["Label_2"][i].text()))

        self.plot_list_4_14 = [int(self.labelResult1_4_14.text().count("True")),
                               int(self.labelResult2_4_14.text().count("True"))]

    def table4_15(self):  # neo tab 2
        self.model = TableModel(self.artistic)
        self.table_view_21.setModel(self.model)

        self.table_view_21.verticalHeader().setVisible(False)
        self.table_view_21.horizontalHeader().setVisible(False)
        self.table_view_21.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_21.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_15_1, self.check4_15_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.artistic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_29 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_29)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_15_1.append(self.chkBoxItem4_29)
            self.table_view_21.setIndexWidget(self.table_view_21.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.artistic.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_30 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_30)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_15_2.append(self.chkBoxItem4_30)
            self.table_view_21.setIndexWidget(self.table_view_21.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_15 = pd.DataFrame({"1": self.check4_15_1, "2": self.check4_15_2})
        a = []
        for i in range(len(self.artistic.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_15 = pd.DataFrame({"Label_1": a})
        self.df4_15["Label_2"] = self.df4_15["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_15(self):
        self.labelResult1_4_15.setText("")
        for i, v in enumerate(self.checkBox_data_4_15["1"]):
            self.df4_15["Label_1"][i] = QLabel()
            self.df4_15["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_15.setText("{}, {}".format(self.labelResult1_4_15.text(),
                                                           self.df4_15["Label_1"][i].text()))
        self.labelResult2_4_15.setText("")
        for i, v in enumerate(self.checkBox_data_4_15["2"]):
            self.df4_15["Label_2"][i] = QLabel()
            self.df4_15["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_15.setText("{}, {}".format(self.labelResult2_4_15.text(),
                                                           self.df4_15["Label_2"][i].text()))

        self.plot_list_4_15 = [int(self.labelResult1_4_15.text().count("True")),
                               int(self.labelResult2_4_15.text().count("True"))]

    def table4_16(self):  # neo tab 2
        self.model = TableModel(self.social)
        self.table_view_22.setModel(self.model)

        self.table_view_22.verticalHeader().setVisible(False)
        self.table_view_22.horizontalHeader().setVisible(False)
        self.table_view_22.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_22.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_16_1, self.check4_16_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.social.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_31 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_31)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_16_1.append(self.chkBoxItem4_31)
            self.table_view_22.setIndexWidget(self.table_view_22.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.social.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_32 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_32)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_16_2.append(self.chkBoxItem4_32)
            self.table_view_22.setIndexWidget(self.table_view_22.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_16 = pd.DataFrame({"1": self.check4_16_1, "2": self.check4_16_2})
        a = []
        for i in range(len(self.social.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_16 = pd.DataFrame({"Label_1": a})
        self.df4_16["Label_2"] = self.df4_16["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_16(self):
        self.labelResult1_4_16.setText("")
        for i, v in enumerate(self.checkBox_data_4_16["1"]):
            self.df4_16["Label_1"][i] = QLabel()
            self.df4_16["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_16.setText("{}, {}".format(self.labelResult1_4_16.text(),
                                                           self.df4_16["Label_1"][i].text()))
        self.labelResult2_4_16.setText("")
        for i, v in enumerate(self.checkBox_data_4_16["2"]):
            self.df4_16["Label_2"][i] = QLabel()
            self.df4_16["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_16.setText("{}, {}".format(self.labelResult2_4_16.text(),
                                                           self.df4_16["Label_2"][i].text()))

        self.plot_list_4_16 = [int(self.labelResult1_4_16.text().count("True")),
                               int(self.labelResult2_4_16.text().count("True"))]

    def table4_17(self):  # neo tab 2
        self.model = TableModel(self.entreprising)
        self.table_view_23.setModel(self.model)

        self.table_view_23.verticalHeader().setVisible(False)
        self.table_view_23.horizontalHeader().setVisible(False)
        self.table_view_23.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_23.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_17_1, self.check4_17_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.entreprising.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_33 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_33)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_17_1.append(self.chkBoxItem4_33)
            self.table_view_23.setIndexWidget(self.table_view_23.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.entreprising.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_34 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_34)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_17_2.append(self.chkBoxItem4_34)
            self.table_view_23.setIndexWidget(self.table_view_23.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_17 = pd.DataFrame({"1": self.check4_17_1, "2": self.check4_17_2})
        a = []
        for i in range(len(self.entreprising.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_17 = pd.DataFrame({"Label_1": a})
        self.df4_17["Label_2"] = self.df4_17["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_17(self):
        self.labelResult1_4_17.setText("")
        for i, v in enumerate(self.checkBox_data_4_17["1"]):
            self.df4_17["Label_1"][i] = QLabel()
            self.df4_17["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_17.setText("{}, {}".format(self.labelResult1_4_17.text(),
                                                           self.df4_17["Label_1"][i].text()))
        self.labelResult2_4_17.setText("")
        for i, v in enumerate(self.checkBox_data_4_17["2"]):
            self.df4_17["Label_2"][i] = QLabel()
            self.df4_17["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_17.setText("{}, {}".format(self.labelResult2_4_17.text(),
                                                           self.df4_17["Label_2"][i].text()))

        self.plot_list_4_17 = [int(self.labelResult1_4_17.text().count("True")),
                               int(self.labelResult2_4_17.text().count("True"))]

    def table4_18(self):  # neo tab 2
        self.model = TableModel(self.conventional)
        self.table_view_24.setModel(self.model)

        self.table_view_24.verticalHeader().setVisible(False)
        self.table_view_24.horizontalHeader().setVisible(False)
        self.table_view_24.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_24.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_18_1, self.check4_18_2 = [], []
        # self.check3=[]
        for i, v in enumerate(self.conventional.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_35 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_35)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_18_1.append(self.chkBoxItem4_35)
            self.table_view_24.setIndexWidget(self.table_view_24.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.conventional.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_36 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_36)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_18_2.append(self.chkBoxItem4_36)
            self.table_view_24.setIndexWidget(self.table_view_24.model().index(i, 2), self.centerdCheckBoxWidget)

        self.checkBox_data_4_18 = pd.DataFrame({"1": self.check4_18_1, "2": self.check4_18_2})
        a = []
        for i in range(len(self.conventional.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_18 = pd.DataFrame({"Label_1": a})
        self.df4_18["Label_2"] = self.df4_18["Label_1"]

        # print(self.checkBox_data_3)

    def checkboxChanged_4_18(self):
        self.labelResult1_4_18.setText("")
        for i, v in enumerate(self.checkBox_data_4_18["1"]):
            self.df4_18["Label_1"][i] = QLabel()
            self.df4_18["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_18.setText("{}, {}".format(self.labelResult1_4_18.text(),
                                                           self.df4_18["Label_1"][i].text()))
        self.labelResult2_4_18.setText("")
        for i, v in enumerate(self.checkBox_data_4_18["2"]):
            self.df4_18["Label_2"][i] = QLabel()
            self.df4_18["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_18.setText("{}, {}".format(self.labelResult2_4_18.text(),
                                                           self.df4_18["Label_2"][i].text()))

        self.plot_list_4_18 = [int(self.labelResult1_4_18.text().count("True")),
                               int(self.labelResult2_4_18.text().count("True"))]

    def table4_19(self):  # neo tab 2
        self.model = TableModel(self.self_assessment__1)
        self.table_view_4.setModel(self.model)

        self.table_view_4.verticalHeader().setVisible(False)
        self.table_view_4.horizontalHeader().setVisible(False)
        self.table_view_4.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_4.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_19_1, self.check4_19_2, self.check4_19_3, self.check4_19_4, self.check4_19_5, self.check4_19_6 = [], [], [], [], [], []
        # self.check3=[]
        for i, v in enumerate(self.self_assessment__1.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_37 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_37)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_19_1.append(self.chkBoxItem4_37)
            self.table_view_4.setIndexWidget(self.table_view_4.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.self_assessment__1.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_38 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_38)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_19_2.append(self.chkBoxItem4_38)
            self.table_view_4.setIndexWidget(self.table_view_4.model().index(i, 2), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.self_assessment__1.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_39 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_39)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_19_3.append(self.chkBoxItem4_39)
            self.table_view_4.setIndexWidget(self.table_view_4.model().index(i, 3), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.self_assessment__1.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_40 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_40)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_19_4.append(self.chkBoxItem4_40)
            self.table_view_4.setIndexWidget(self.table_view_4.model().index(i, 4), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.self_assessment__1.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_41 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_41)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_19_5.append(self.chkBoxItem4_41)
            self.table_view_4.setIndexWidget(self.table_view_4.model().index(i, 5), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.self_assessment__1.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_42 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_42)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_19_6.append(self.chkBoxItem4_42)
            self.table_view_4.setIndexWidget(self.table_view_4.model().index(i, 6), self.centerdCheckBoxWidget)

        self.checkBox_data_4_19 = pd.DataFrame({"1": self.check4_19_1, "2": self.check4_19_2,
                                                "3": self.check4_19_3, "4": self.check4_19_4,
                                                "5": self.check4_19_5, "6": self.check4_19_6})
        a = []
        for i in range(len(self.self_assessment__1.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_19 = pd.DataFrame({"Label_1": a})
        self.df4_19["Label_2"] = self.df4_19["Label_1"]
        self.df4_19["Label_3"] = self.df4_19["Label_1"]
        self.df4_19["Label_4"] = self.df4_19["Label_1"]
        self.df4_19["Label_5"] = self.df4_19["Label_1"]
        self.df4_19["Label_6"] = self.df4_19["Label_1"]
        # print(self.checkBox_data_3)

    def checkboxChanged_4_19(self):
        self.labelResult1_4_19.setText("")
        for i, v in enumerate(self.checkBox_data_4_19["1"]):
            self.df4_19["Label_1"][i] = QLabel()
            self.df4_19["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_19.setText("{}, {}".format(self.labelResult1_4_19.text(),
                                                           self.df4_19["Label_1"][i].text()))
        self.labelResult2_4_19.setText("")
        for i, v in enumerate(self.checkBox_data_4_19["2"]):
            self.df4_19["Label_2"][i] = QLabel()
            self.df4_19["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_19.setText("{}, {}".format(self.labelResult2_4_19.text(),
                                                           self.df4_19["Label_2"][i].text()))
        self.labelResult3_4_19.setText("")
        for i, v in enumerate(self.checkBox_data_4_19["3"]):
            self.df4_19["Label_3"][i] = QLabel()
            self.df4_19["Label_3"][i].setText("True" if v.checkState() else "False")
            self.labelResult3_4_19.setText("{}, {}".format(self.labelResult3_4_19.text(),
                                                           self.df4_19["Label_3"][i].text()))
        self.labelResult4_4_19.setText("")
        for i, v in enumerate(self.checkBox_data_4_19["4"]):
            self.df4_19["Label_4"][i] = QLabel()
            self.df4_19["Label_4"][i].setText("True" if v.checkState() else "False")
            self.labelResult4_4_19.setText("{}, {}".format(self.labelResult4_4_19.text(),
                                                           self.df4_19["Label_4"][i].text()))
        self.labelResult5_4_19.setText("")
        for i, v in enumerate(self.checkBox_data_4_19["5"]):
            self.df4_19["Label_5"][i] = QLabel()
            self.df4_19["Label_5"][i].setText("True" if v.checkState() else "False")
            self.labelResult5_4_19.setText("{}, {}".format(self.labelResult5_4_19.text(),
                                                           self.df4_19["Label_5"][i].text()))
        self.labelResult6_4_19.setText("")
        for i, v in enumerate(self.checkBox_data_4_19["6"]):
            self.df4_19["Label_6"][i] = QLabel()
            self.df4_19["Label_6"][i].setText("True" if v.checkState() else "False")
            self.labelResult6_4_19.setText("{}, {}".format(self.labelResult6_4_19.text(),
                                                           self.df4_19["Label_6"][i].text()))

        self.plot_list_4_19 = [int(self.labelResult1_4_19.text().count("True")),
                               int(self.labelResult2_4_19.text().count("True")),
                               int(self.labelResult3_4_19.text().count("True")),
                               int(self.labelResult4_4_19.text().count("True")),
                               int(self.labelResult5_4_19.text().count("True")),
                               int(self.labelResult6_4_19.text().count("True"))]

    def table4_20(self):  # neo tab 2
        self.model = TableModel(self.self_assessment__2)
        self.table_view_5.setModel(self.model)

        self.table_view_5.verticalHeader().setVisible(False)
        self.table_view_5.horizontalHeader().setVisible(False)
        self.table_view_5.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_5.resizeColumnsToContents()
        # self.table_view_4.resizeRowsToContents()

        self.check4_20_1, self.check4_20_2, self.check4_20_3, self.check4_20_4, self.check4_20_5, self.check4_20_6 = [], [], [], [], [], []
        # self.check3=[]
        for i, v in enumerate(self.self_assessment__2.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_43 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_43)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_20_1.append(self.chkBoxItem4_43)
            self.table_view_5.setIndexWidget(self.table_view_5.model().index(i, 1), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.self_assessment__2.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_44 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_44)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_20_2.append(self.chkBoxItem4_44)
            self.table_view_5.setIndexWidget(self.table_view_5.model().index(i, 2), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.self_assessment__2.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_45 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_45)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_20_3.append(self.chkBoxItem4_45)
            self.table_view_5.setIndexWidget(self.table_view_5.model().index(i, 3), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.self_assessment__2.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_46 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_46)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_20_4.append(self.chkBoxItem4_46)
            self.table_view_5.setIndexWidget(self.table_view_5.model().index(i, 4), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.self_assessment__2.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_47 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_47)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_20_5.append(self.chkBoxItem4_47)
            self.table_view_5.setIndexWidget(self.table_view_5.model().index(i, 5), self.centerdCheckBoxWidget)
        for i, v in enumerate(self.self_assessment__2.iloc[1:, -1], start=1):
            self.centerdCheckBoxWidget = QWidget()
            self.chkBoxItem4_48 = QCheckBox(str(v))
            self.checkBoxLayout = QHBoxLayout(self.centerdCheckBoxWidget)
            self.checkBoxLayout.addWidget(self.chkBoxItem4_48)
            self.checkBoxLayout.setAlignment(Qt.AlignCenter)
            self.checkBoxLayout.setContentsMargins(0, 0, 0, 0)
            self.centerdCheckBoxWidget.setLayout(self.checkBoxLayout)
            self.check4_20_6.append(self.chkBoxItem4_48)
            self.table_view_5.setIndexWidget(self.table_view_5.model().index(i, 6), self.centerdCheckBoxWidget)

        self.checkBox_data_4_20 = pd.DataFrame({"1": self.check4_20_1, "2": self.check4_20_2,
                                                "3": self.check4_20_3, "4": self.check4_20_4,
                                                "5": self.check4_20_5, "6": self.check4_20_6})
        a = []
        for i in range(len(self.self_assessment__2.iloc[:, -1])):
            i = ""
            a.append(i)
        self.df4_20 = pd.DataFrame({"Label_1": a})
        self.df4_20["Label_2"] = self.df4_20["Label_1"]
        self.df4_20["Label_3"] = self.df4_20["Label_1"]
        self.df4_20["Label_4"] = self.df4_20["Label_1"]
        self.df4_20["Label_5"] = self.df4_20["Label_1"]
        self.df4_20["Label_6"] = self.df4_20["Label_1"]
        # print(self.checkBox_data_3)

    def checkboxChanged_4_20(self):
        self.labelResult1_4_20.setText("")
        for i, v in enumerate(self.checkBox_data_4_20["1"]):
            self.df4_20["Label_1"][i] = QLabel()
            self.df4_20["Label_1"][i].setText("True" if v.checkState() else "False")
            self.labelResult1_4_20.setText("{}, {}".format(self.labelResult1_4_20.text(),
                                                           self.df4_20["Label_1"][i].text()))
        self.labelResult2_4_20.setText("")
        for i, v in enumerate(self.checkBox_data_4_20["2"]):
            self.df4_20["Label_2"][i] = QLabel()
            self.df4_20["Label_2"][i].setText("True" if v.checkState() else "False")
            self.labelResult2_4_20.setText("{}, {}".format(self.labelResult2_4_20.text(),
                                                           self.df4_20["Label_2"][i].text()))
        self.labelResult3_4_20.setText("")
        for i, v in enumerate(self.checkBox_data_4_20["3"]):
            self.df4_20["Label_3"][i] = QLabel()
            self.df4_20["Label_3"][i].setText("True" if v.checkState() else "False")
            self.labelResult3_4_20.setText("{}, {}".format(self.labelResult3_4_20.text(),
                                                           self.df4_20["Label_3"][i].text()))
        self.labelResult4_4_20.setText("")
        for i, v in enumerate(self.checkBox_data_4_20["4"]):
            self.df4_20["Label_4"][i] = QLabel()
            self.df4_20["Label_4"][i].setText("True" if v.checkState() else "False")
            self.labelResult4_4_20.setText("{}, {}".format(self.labelResult4_4_20.text(),
                                                           self.df4_20["Label_4"][i].text()))
        self.labelResult5_4_20.setText("")
        for i, v in enumerate(self.checkBox_data_4_20["5"]):
            self.df4_20["Label_5"][i] = QLabel()
            self.df4_20["Label_5"][i].setText("True" if v.checkState() else "False")
            self.labelResult5_4_20.setText("{}, {}".format(self.labelResult5_4_20.text(),
                                                           self.df4_20["Label_5"][i].text()))
        self.labelResult6_4_20.setText("")
        for i, v in enumerate(self.checkBox_data_4_20["6"]):
            self.df4_20["Label_6"][i] = QLabel()
            self.df4_20["Label_6"][i].setText("True" if v.checkState() else "False")
            self.labelResult6_4_20.setText("{}, {}".format(self.labelResult6_4_20.text(),
                                                           self.df4_20["Label_6"][i].text()))

        self.plot_list_4_20 = [int(self.labelResult1_4_20.text().count("True")),
                               int(self.labelResult2_4_20.text().count("True")),
                               int(self.labelResult3_4_20.text().count("True")),
                               int(self.labelResult4_4_20.text().count("True")),
                               int(self.labelResult5_4_20.text().count("True")),
                               int(self.labelResult6_4_20.text().count("True"))]

    def table5_1(self):
        self.model = TableModel(self.que_1)
        self.table_view_25.setModel(self.model)

        self.table_view_25.verticalHeader().setVisible(False)
        self.table_view_25.horizontalHeader().setVisible(False)
        self.table_view_25.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_25.resizeColumnsToContents()
        # self.table_view.resizeRowsToContents()

        # self.check5_1=[]
        self.centerdCheckBoxWidget_5_1 = QWidget()
        self.chkBoxItem5_1_1 = QCheckBox()
        self.checkBoxLayout_5_1 = QHBoxLayout(self.centerdCheckBoxWidget_5_1)
        self.checkBoxLayout_5_1.addWidget(self.chkBoxItem5_1_1)
        self.checkBoxLayout_5_1.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_1.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_1.setLayout(self.checkBoxLayout_5_1)
        self.table_view_25.setIndexWidget(self.table_view_25.model().index(0, 1), self.centerdCheckBoxWidget_5_1)

        self.centerdCheckBoxWidget_5_1_2 = QWidget()
        self.chkBoxItem5_1_2 = QCheckBox()
        self.checkBoxLayout_5_1_2 = QHBoxLayout(self.centerdCheckBoxWidget_5_1_2)
        self.checkBoxLayout_5_1_2.addWidget(self.chkBoxItem5_1_2)
        self.checkBoxLayout_5_1_2.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_1_2.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_1_2.setLayout(self.checkBoxLayout_5_1_2)
        self.table_view_25.setIndexWidget(self.table_view_25.model().index(1, 1), self.centerdCheckBoxWidget_5_1_2)

    def table5_2(self):
        self.model = TableModel(self.que_2)
        self.table_view_52.setModel(self.model)

        self.table_view_52.verticalHeader().setVisible(False)
        self.table_view_52.horizontalHeader().setVisible(False)
        self.table_view_52.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_52.resizeColumnsToContents()
        # self.table_view.resizeRowsToContents()

        self.centerdCheckBoxWidget_5_2 = QWidget()
        self.chkBoxItem5_2 = QCheckBox()
        self.checkBoxLayout_5_2 = QHBoxLayout(self.centerdCheckBoxWidget_5_2)
        self.checkBoxLayout_5_2.addWidget(self.chkBoxItem5_2)
        self.checkBoxLayout_5_2.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_2.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_2.setLayout(self.checkBoxLayout_5_2)
        self.table_view_52.setIndexWidget(self.table_view_52.model().index(0, 1), self.centerdCheckBoxWidget_5_2)

    def table5_3(self):
        self.model = TableModel(self.que_drastiriotites)
        self.table_view_53.setModel(self.model)

        self.table_view_53.verticalHeader().setVisible(False)
        self.table_view_53.horizontalHeader().setVisible(False)
        self.table_view_53.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_53.resizeColumnsToContents()
        # self.table_view.resizeRowsToContents()

        self.centerdCheckBoxWidget_5_3_1 = QWidget()
        self.chkBoxItem5_3_1 = QCheckBox()
        self.checkBoxLayout_5_3_1 = QHBoxLayout(self.centerdCheckBoxWidget_5_3_1)
        self.checkBoxLayout_5_3_1.addWidget(self.chkBoxItem5_3_1)
        self.checkBoxLayout_5_3_1.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_3_1.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_3_1.setLayout(self.checkBoxLayout_5_3_1)
        self.table_view_53.setIndexWidget(self.table_view_53.model().index(0, 1), self.centerdCheckBoxWidget_5_3_1)

        self.centerdCheckBoxWidget_5_3_2 = QWidget()
        self.chkBoxItem5_3_2 = QCheckBox()
        self.checkBoxLayout_5_3_2 = QHBoxLayout(self.centerdCheckBoxWidget_5_3_2)
        self.checkBoxLayout_5_3_2.addWidget(self.chkBoxItem5_3_2)
        self.checkBoxLayout_5_3_2.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_3_2.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_3_2.setLayout(self.checkBoxLayout_5_3_2)
        self.table_view_53.setIndexWidget(self.table_view_53.model().index(1, 1), self.centerdCheckBoxWidget_5_3_2)

        self.centerdCheckBoxWidget_5_3_3 = QWidget()
        self.chkBoxItem5_3_3 = QCheckBox()
        self.checkBoxLayout_5_3_3 = QHBoxLayout(self.centerdCheckBoxWidget_5_3_3)
        self.checkBoxLayout_5_3_3.addWidget(self.chkBoxItem5_3_3)
        self.checkBoxLayout_5_3_3.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_3_3.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_3_3.setLayout(self.checkBoxLayout_5_3_3)
        self.table_view_53.setIndexWidget(self.table_view_53.model().index(2, 1), self.centerdCheckBoxWidget_5_3_3)

        self.centerdCheckBoxWidget_5_3_4 = QWidget()
        self.chkBoxItem5_3_4 = QCheckBox()
        self.checkBoxLayout_5_3_4 = QHBoxLayout(self.centerdCheckBoxWidget_5_3_4)
        self.checkBoxLayout_5_3_4.addWidget(self.chkBoxItem5_3_4)
        self.checkBoxLayout_5_3_4.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_3_4.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_3_4.setLayout(self.checkBoxLayout_5_3_4)
        self.table_view_53.setIndexWidget(self.table_view_53.model().index(3, 1), self.centerdCheckBoxWidget_5_3_4)

        self.centerdCheckBoxWidget_5_3_5 = QWidget()
        self.chkBoxItem5_3_5 = QCheckBox()
        self.checkBoxLayout_5_3_5 = QHBoxLayout(self.centerdCheckBoxWidget_5_3_5)
        self.checkBoxLayout_5_3_5.addWidget(self.chkBoxItem5_3_5)
        self.checkBoxLayout_5_3_5.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_3_5.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_3_5.setLayout(self.checkBoxLayout_5_3_5)
        self.table_view_53.setIndexWidget(self.table_view_53.model().index(4, 1), self.centerdCheckBoxWidget_5_3_5)

        self.centerdCheckBoxWidget_5_3_6 = QWidget()
        self.chkBoxItem5_3_6 = QCheckBox()
        self.checkBoxLayout_5_3_6 = QHBoxLayout(self.centerdCheckBoxWidget_5_3_6)
        self.checkBoxLayout_5_3_6.addWidget(self.chkBoxItem5_3_6)
        self.checkBoxLayout_5_3_6.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_3_6.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_3_6.setLayout(self.checkBoxLayout_5_3_6)
        self.table_view_53.setIndexWidget(self.table_view_53.model().index(5, 1), self.centerdCheckBoxWidget_5_3_6)

    def table5_4(self):
        self.model = TableModel(self.que_ikanotites)
        self.table_view_54.setModel(self.model)

        self.table_view_54.verticalHeader().setVisible(False)
        self.table_view_54.horizontalHeader().setVisible(False)
        self.table_view_54.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_54.resizeColumnsToContents()
        # self.table_view.resizeRowsToContents()

        self.centerdCheckBoxWidget_5_4_1 = QWidget()
        self.chkBoxItem5_4_1 = QCheckBox()
        self.checkBoxLayout_5_4_1 = QHBoxLayout(self.centerdCheckBoxWidget_5_4_1)
        self.checkBoxLayout_5_4_1.addWidget(self.chkBoxItem5_4_1)
        self.checkBoxLayout_5_4_1.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_4_1.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_4_1.setLayout(self.checkBoxLayout_5_4_1)
        self.table_view_54.setIndexWidget(self.table_view_54.model().index(0, 1), self.centerdCheckBoxWidget_5_4_1)

        self.centerdCheckBoxWidget_5_4_2 = QWidget()
        self.chkBoxItem5_4_2 = QCheckBox()
        self.checkBoxLayout_5_4_2 = QHBoxLayout(self.centerdCheckBoxWidget_5_4_2)
        self.checkBoxLayout_5_4_2.addWidget(self.chkBoxItem5_4_2)
        self.checkBoxLayout_5_4_2.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_4_2.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_4_2.setLayout(self.checkBoxLayout_5_4_2)
        self.table_view_54.setIndexWidget(self.table_view_54.model().index(1, 1), self.centerdCheckBoxWidget_5_4_2)

        self.centerdCheckBoxWidget_5_4_3 = QWidget()
        self.chkBoxItem5_4_3 = QCheckBox()
        self.checkBoxLayout_5_4_3 = QHBoxLayout(self.centerdCheckBoxWidget_5_4_3)
        self.checkBoxLayout_5_4_3.addWidget(self.chkBoxItem5_4_3)
        self.checkBoxLayout_5_4_3.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_4_3.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_4_3.setLayout(self.checkBoxLayout_5_4_3)
        self.table_view_54.setIndexWidget(self.table_view_54.model().index(2, 1), self.centerdCheckBoxWidget_5_4_3)

        self.centerdCheckBoxWidget_5_4_4 = QWidget()
        self.chkBoxItem5_4_4 = QCheckBox()
        self.checkBoxLayout_5_4_4 = QHBoxLayout(self.centerdCheckBoxWidget_5_4_4)
        self.checkBoxLayout_5_4_4.addWidget(self.chkBoxItem5_4_4)
        self.checkBoxLayout_5_4_4.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_4_4.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_4_4.setLayout(self.checkBoxLayout_5_4_4)
        self.table_view_54.setIndexWidget(self.table_view_54.model().index(3, 1), self.centerdCheckBoxWidget_5_4_4)

        self.centerdCheckBoxWidget_5_4_5 = QWidget()
        self.chkBoxItem5_4_5 = QCheckBox()
        self.checkBoxLayout_5_4_5 = QHBoxLayout(self.centerdCheckBoxWidget_5_4_5)
        self.checkBoxLayout_5_4_5.addWidget(self.chkBoxItem5_4_5)
        self.checkBoxLayout_5_4_5.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_4_5.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_4_5.setLayout(self.checkBoxLayout_5_4_5)
        self.table_view_54.setIndexWidget(self.table_view_54.model().index(4, 1), self.centerdCheckBoxWidget_5_4_5)

        self.centerdCheckBoxWidget_5_4_6 = QWidget()
        self.chkBoxItem5_4_6 = QCheckBox()
        self.checkBoxLayout_5_4_6 = QHBoxLayout(self.centerdCheckBoxWidget_5_4_6)
        self.checkBoxLayout_5_4_6.addWidget(self.chkBoxItem5_4_6)
        self.checkBoxLayout_5_4_6.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_4_6.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_4_6.setLayout(self.checkBoxLayout_5_4_6)
        self.table_view_54.setIndexWidget(self.table_view_54.model().index(5, 1), self.centerdCheckBoxWidget_5_4_6)

    def table5_5(self):
        self.model = TableModel(self.que_epaggelma)
        self.table_view_55.setModel(self.model)

        self.table_view_55.verticalHeader().setVisible(False)
        self.table_view_55.horizontalHeader().setVisible(False)
        self.table_view_55.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_55.resizeColumnsToContents()
        # self.table_view.resizeRowsToContents()

        self.centerdCheckBoxWidget_5_5_1 = QWidget()
        self.chkBoxItem5_5_1 = QCheckBox()
        self.checkBoxLayout_5_5_1 = QHBoxLayout(self.centerdCheckBoxWidget_5_5_1)
        self.checkBoxLayout_5_5_1.addWidget(self.chkBoxItem5_5_1)
        self.checkBoxLayout_5_5_1.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_5_1.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_5_1.setLayout(self.checkBoxLayout_5_5_1)
        self.table_view_55.setIndexWidget(self.table_view_55.model().index(0, 1), self.centerdCheckBoxWidget_5_5_1)

        self.centerdCheckBoxWidget_5_5_2 = QWidget()
        self.chkBoxItem5_5_2 = QCheckBox()
        self.checkBoxLayout_5_5_2 = QHBoxLayout(self.centerdCheckBoxWidget_5_5_2)
        self.checkBoxLayout_5_5_2.addWidget(self.chkBoxItem5_5_2)
        self.checkBoxLayout_5_5_2.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_5_2.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_5_2.setLayout(self.checkBoxLayout_5_5_2)
        self.table_view_55.setIndexWidget(self.table_view_55.model().index(1, 1), self.centerdCheckBoxWidget_5_5_2)

        self.centerdCheckBoxWidget_5_5_3 = QWidget()
        self.chkBoxItem5_5_3 = QCheckBox()
        self.checkBoxLayout_5_5_3 = QHBoxLayout(self.centerdCheckBoxWidget_5_5_3)
        self.checkBoxLayout_5_5_3.addWidget(self.chkBoxItem5_5_3)
        self.checkBoxLayout_5_5_3.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_5_3.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_5_3.setLayout(self.checkBoxLayout_5_5_3)
        self.table_view_55.setIndexWidget(self.table_view_55.model().index(2, 1), self.centerdCheckBoxWidget_5_5_3)

        self.centerdCheckBoxWidget_5_5_4 = QWidget()
        self.chkBoxItem5_5_4 = QCheckBox()
        self.checkBoxLayout_5_5_4 = QHBoxLayout(self.centerdCheckBoxWidget_5_5_4)
        self.checkBoxLayout_5_5_4.addWidget(self.chkBoxItem5_5_4)
        self.checkBoxLayout_5_5_4.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_5_4.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_5_4.setLayout(self.checkBoxLayout_5_5_4)
        self.table_view_55.setIndexWidget(self.table_view_55.model().index(3, 1), self.centerdCheckBoxWidget_5_5_4)

        self.centerdCheckBoxWidget_5_5_5 = QWidget()
        self.chkBoxItem5_5_5 = QCheckBox()
        self.checkBoxLayout_5_5_5 = QHBoxLayout(self.centerdCheckBoxWidget_5_5_5)
        self.checkBoxLayout_5_5_5.addWidget(self.chkBoxItem5_5_5)
        self.checkBoxLayout_5_5_5.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_5_5.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_5_5.setLayout(self.checkBoxLayout_5_5_5)
        self.table_view_55.setIndexWidget(self.table_view_55.model().index(4, 1), self.centerdCheckBoxWidget_5_5_5)

        self.centerdCheckBoxWidget_5_5_6 = QWidget()
        self.chkBoxItem5_5_6 = QCheckBox()
        self.checkBoxLayout_5_5_6 = QHBoxLayout(self.centerdCheckBoxWidget_5_5_6)
        self.checkBoxLayout_5_5_6.addWidget(self.chkBoxItem5_5_6)
        self.checkBoxLayout_5_5_6.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_5_5_6.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_5_5_6.setLayout(self.checkBoxLayout_5_5_6)
        self.table_view_55.setIndexWidget(self.table_view_55.model().index(5, 1), self.centerdCheckBoxWidget_5_5_6)

    def table5_6(self):
        self.model = TableModel(self.que_aksiologisi)
        self.table_view_56.setModel(self.model)

        self.table_view_56.verticalHeader().setVisible(False)
        self.table_view_56.horizontalHeader().setVisible(False)
        self.table_view_56.setSizeAdjustPolicy(
            QtWidgets.QAbstractScrollArea.AdjustToContents)
        self.table_view_56.resizeColumnsToContents()
        # self.table_view.resizeRowsToContents()

        self.centerdCheckBoxWidget_6_1_1 = QWidget()
        self.chkBoxItem6_1_1 = QCheckBox()
        self.checkBoxLayout_6_1_1 = QHBoxLayout(self.centerdCheckBoxWidget_6_1_1)
        self.checkBoxLayout_6_1_1.addWidget(self.chkBoxItem6_1_1)
        self.checkBoxLayout_6_1_1.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_6_1_1.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_6_1_1.setLayout(self.checkBoxLayout_6_1_1)
        self.table_view_56.setIndexWidget(self.table_view_56.model().index(0, 1), self.centerdCheckBoxWidget_6_1_1)

        self.centerdCheckBoxWidget_6_1_2 = QWidget()
        self.chkBoxItem6_1_2 = QCheckBox()
        self.checkBoxLayout_6_1_2 = QHBoxLayout(self.centerdCheckBoxWidget_6_1_2)
        self.checkBoxLayout_6_1_2.addWidget(self.chkBoxItem6_1_2)
        self.checkBoxLayout_6_1_2.setAlignment(Qt.AlignCenter)
        self.checkBoxLayout_6_1_2.setContentsMargins(0, 0, 0, 0)
        self.centerdCheckBoxWidget_6_1_1.setLayout(self.checkBoxLayout_6_1_2)
        self.table_view_56.setIndexWidget(self.table_view_56.model().index(1, 1), self.centerdCheckBoxWidget_6_1_2)

    def R(self):
        R0 = (int(self.labelResult1_4_1.text().count("True")) + int(self.labelResult1_4_7.text().count("True")) +
              int(self.labelResult1_4_13.text().count("True")))

        a = self.labelResult1_4_19.text()
        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        for i in range(len(self.self_assessment__1.iloc[:, -1]) - 1):
            if x[i] == 1:
                y_R = 7 - i

        b = self.labelResult1_4_20.text()
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        for i in range(len(self.self_assessment__2.iloc[:, -1]) - 1):
            if y[i] == 1:
                z_R = 7 - i
                # self.R
        # print(type(R0 + y_R + z_R))
        print("R=", R0 + y_R + z_R)

        return R0 + y_R + z_R

    def I(self):
        I0 = (int(self.labelResult1_4_2.text().count("True")) + int(self.labelResult1_4_8.text().count("True")) +
              int(self.labelResult1_4_14.text().count("True")))

        a = self.labelResult2_4_19.text()
        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        for i in range(len(self.self_assessment__1.iloc[:, -1]) - 1):
            if x[i] == 1:
                y_I = 7 - i
        b = self.labelResult2_4_20.text()
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        for i in range(len(self.self_assessment__2.iloc[:, -1]) - 1):
            if y[i] == 1:
                z_I = 7 - i
                # self.R
        print("I=", I0 + y_I + z_I)
        return I0 + y_I + z_I

    def A(self):
        A0 = (int(self.labelResult1_4_3.text().count("True")) + int(self.labelResult1_4_9.text().count("True")) +
              int(self.labelResult1_4_15.text().count("True")))

        a = self.labelResult3_4_19.text()
        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        for i in range(len(self.self_assessment__1.iloc[:, -1]) - 1):
            if x[i] == 1:
                y_A = 7 - i
        b = self.labelResult3_4_20.text()
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        for i in range(len(self.self_assessment__2.iloc[:, -1]) - 1):
            if y[i] == 1:
                z_A = 7 - i
                # self.R
        print("A=", A0 + y_A + z_A)
        return A0 + y_A + z_A

    def S(self):
        S0 = (int(self.labelResult1_4_4.text().count("True")) + int(self.labelResult1_4_10.text().count("True")) +
              int(self.labelResult1_4_16.text().count("True")))

        a = self.labelResult4_4_19.text()
        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        for i in range(len(self.self_assessment__1.iloc[:, -1]) - 1):
            if x[i] == 1:
                y_S = 7 - i
        b = self.labelResult4_4_20.text()
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        for i in range(len(self.self_assessment__2.iloc[:, -1]) - 1):
            if y[i] == 1:
                z_S = 7 - i
                # self.R
        print("S=", S0 + y_S + z_S)
        return S0 + y_S + z_S

    def E(self):
        E0 = (int(self.labelResult1_4_5.text().count("True")) + int(self.labelResult1_4_11.text().count("True")) +
              int(self.labelResult1_4_17.text().count("True")))

        a = self.labelResult5_4_19.text()
        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        for i in range(len(self.self_assessment__1.iloc[:, -1]) - 1):
            if x[i] == 1:
                y_E = 7 - i
        b = self.labelResult5_4_20.text()
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        for i in range(len(self.self_assessment__2.iloc[:, -1]) - 1):
            if y[i] == 1:
                z_E = 7 - i
                # self.R
        print("E=", E0 + y_E + z_E)
        return E0 + y_E + z_E

    def C(self):
        C0 = (int(self.labelResult1_4_6.text().count("True")) + int(self.labelResult1_4_12.text().count("True")) +
              int(self.labelResult1_4_18.text().count("True")))

        a = self.labelResult6_4_19.text()
        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        for i in range(len(self.self_assessment__1.iloc[:, -1]) - 1):
            if x[i] == 1:
                y_C = 7 - i
        b = self.labelResult6_4_20.text()
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        for i in range(len(self.self_assessment__2.iloc[:, -1]) - 1):
            if y[i] == 1:
                z_C = 7 - i
                # self.R
        print("C=", C0 + y_C + z_C)
        return C0 + y_C + z_C

    # def df_hol(self):
    #     return pd.DataFrame({"R":self.R(),"I":self.I(),"A":self.A(),"S":self.S(),"E":self.E(),"C":self.C(),})

    def hol_results(self):
        self.df_hol = pd.DataFrame(
            {"R": self.R(), "I": self.I(), "A": self.A(), "S": self.S(), "E": self.E(), "C": self.C()}, index=[0])
        print((self.df_hol))
        self.label_25.setText("R = " + str(int(self.df_hol["R"])))
        self.r_write = "Ο ενδιαφερόμενος στην κατηγορία ΠΡΑΚΤΙΚΟΣ (REALISTIC) συγκέντρωσε " + str(
            int(self.df_hol["R"])) + " βαθμούς. "
        self.i_write = "Στην κατηγορία ΕΡΕΥΝΗΤΙΚΟΣ (INVESTIGATIVE) συγκέντρωσε " + str(
            int(self.df_hol["I"])) + " βαθμούς. "
        self.a_b_write = ("Στις κατηγορίες ΚΑΛΛΙΤΕΧΝΙΚΟΣ (ARTISTIC) και ΚΟΙΝΩΝΙΚΟΣ (SOCIAL) συγκέντρωσε " +
                          str(int(self.df_hol["A"])) + " και " + str(int(self.df_hol["S"])) + " αντίστοιχα. ")
        self.e_c_write = (
                    "Τέλος στις κατηγορίες ΕΠΙΧΕΙΡΗΜΑΤΙΚΟΣ (ENTERPRISING) και ΣΥΜΒΑΤΙΚΟΣ (CONVENTIONAL) ο ενδιαφερόμενος συγκέντρωσε " +
                    str(int(self.df_hol["E"])) + " και " + str(int(self.df_hol["C"])) + " αντίστοιχα. ")

        self.textEdit_2.setText(self.r_write + self.i_write + self.a_b_write + self.e_c_write)

    def cross(self, list_labelresult, aksia):
        indices = []
        for idx, value in enumerate(list_labelresult):
            if value == 1:
                indices.append(idx)
        lst = []
        for i in indices:
            lst.append(aksia[i])
        return lst

    def pressed_button1(self):
        a = self.labelResult1_1.text()
        b = self.labelResult2_1.text()
        c = self.labelResult3_1.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        z0 = c.split(",")
        z0.pop(0)
        z = [1 if i == " True" else 0 for i in z0]
        w = []
        for i in range(len(self.data.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1) or (y[i] == z[i] == 1) or (x[i] == z[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == z[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_1_1.setChecked(True)
        else:
            self.label.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_1_1.setChecked(False)

        # if self.label.text()==u"Ολοκληρώθηκε Επιτυχώς"
        # print(self.label.text())

    def pressed_button2(self):
        self.label_2.setText(u"Ολοκληρώθηκε Επιτυχώς")
        if (self.label_2.text() == u"Ολοκληρώθηκε Επιτυχώς"):
            self.chkBoxItem5_1_2.setChecked(True)
        else:
            self.chkBoxItem5_1_2.setChecked(False)

    def pressed_button3(self):

        a = self.labelResult1_3.text()
        b = self.labelResult2_3.text()
        c = self.labelResult3_3.text()
        d = self.labelResult4_3.text()
        e = self.labelResult5_3.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        z0 = c.split(",")
        z0.pop(0)
        z = [1 if i == " True" else 0 for i in z0]
        m0 = d.split(",")
        m0.pop(0)
        m = [1 if i == " True" else 0 for i in m0]
        n0 = e.split(",")
        n0.pop(0)
        n = [1 if i == " True" else 0 for i in n0]

        w = []
        for i in range(len(self.second_data.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1) or (x[i] == z[i] == 1) or (x[i] == m[i] == 1) or (x[i] == n[i] == 1)
                    or (y[i] == z[i] == 1) or (y[i] == m[i] == 1) or (y[i] == n[i] == 1)
                    or (z[i] == m[i] == 1) or (z[i] == n[i] == 1) or (m[i] == n[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == z[i] == m[i] == n[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_3.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_2.setChecked(True)
        else:
            self.label_3.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_2.setChecked(False)

    def pressed_button4(self):
        a = self.labelResult1_4_1.text()
        b = self.labelResult2_4_1.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.drastiriotites_realstic.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_4.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_3_1.setChecked(True)
        else:
            self.label_4.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_3_1.setChecked(False)

    def pressed_button5(self):
        a = self.labelResult1_4_2.text()
        b = self.labelResult2_4_2.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.drastiriotites_investigative.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_5.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_3_2.setChecked(True)
        else:
            self.label_5.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_3_2.setChecked(False)

    def pressed_button6(self):
        a = self.labelResult1_4_3.text()
        b = self.labelResult2_4_3.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.drastiriotites_artistic.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_6.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_3_3.setChecked(True)
        else:
            self.label_6.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_3_3.setChecked(False)

    def pressed_button7(self):
        a = self.labelResult1_4_4.text()
        b = self.labelResult2_4_4.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.drastiriotites_social.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_7.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_3_4.setChecked(True)
        else:
            self.label_7.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_3_4.setChecked(False)

    def pressed_button8(self):
        a = self.labelResult1_4_5.text()
        b = self.labelResult2_4_5.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.drastiriotites_entreprising.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_8.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_3_5.setChecked(True)
        else:
            self.label_8.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_3_5.setChecked(False)

    def pressed_button9(self):
        a = self.labelResult1_4_6.text()
        b = self.labelResult2_4_6.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.drastiriotites_conventional.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_9.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_3_6.setChecked(True)

        else:
            self.label_9.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_3_6.setChecked(False)

    def pressed_button10(self):
        a = self.labelResult1_4_7.text()
        b = self.labelResult2_4_7.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.ikanotites_realstic.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_10.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_4_1.setChecked(True)
        else:
            self.label_10.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_4_1.setChecked(False)

    def pressed_button11(self):
        a = self.labelResult1_4_8.text()
        b = self.labelResult2_4_8.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.ikanotites_investigative.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_11.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_4_2.setChecked(True)

        else:
            self.label_11.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_4_2.setChecked(False)

    def pressed_button12(self):
        a = self.labelResult1_4_9.text()
        b = self.labelResult2_4_9.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.ikanotites_artistic.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_12.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_4_3.setChecked(True)
        else:
            self.label_12.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_4_3.setChecked(False)

    def pressed_button13(self):
        a = self.labelResult1_4_10.text()
        b = self.labelResult2_4_10.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.ikanotites_social.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_13.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_4_4.setChecked(True)
        else:
            self.label_13.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_4_4.setChecked(False)

    def pressed_button14(self):
        a = self.labelResult1_4_11.text()
        b = self.labelResult2_4_11.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.ikanotites_entreprising.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_14.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_4_5.setChecked(True)
        else:
            self.label_14.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_4_5.setChecked(False)

    def pressed_button15(self):
        a = self.labelResult1_4_12.text()
        b = self.labelResult2_4_12.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.ikanotites_conventional.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_15.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_4_6.setChecked(True)
        else:
            self.label_15.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_4_6.setChecked(False)

    def pressed_button16(self):
        a = self.labelResult1_4_13.text()
        b = self.labelResult2_4_13.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.realstic.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_16.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_5_1.setChecked(True)
        else:
            self.label_16.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_5_1.setChecked(False)

    def pressed_button17(self):
        a = self.labelResult1_4_14.text()
        b = self.labelResult2_4_14.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.investigative.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_17.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_5_2.setChecked(True)
        else:
            self.label_17.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_5_2.setChecked(False)

    def pressed_button18(self):
        a = self.labelResult1_4_15.text()
        b = self.labelResult2_4_15.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.artistic.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_18.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_5_3.setChecked(True)
        else:
            self.label_18.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_5_3.setChecked(False)

    def pressed_button19(self):
        a = self.labelResult1_4_16.text()
        b = self.labelResult2_4_16.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.realstic.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_19.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_5_4.setChecked(True)
        else:
            self.label_19.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_5_4.setChecked(False)

    def pressed_button20(self):
        a = self.labelResult1_4_17.text()
        b = self.labelResult2_4_17.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.investigative.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_20.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_5_5.setChecked(True)
        else:
            self.label_20.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_5_5.setChecked(False)

    def pressed_button21(self):
        a = self.labelResult1_4_18.text()
        b = self.labelResult2_4_18.text()

        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]
        w = []
        for i in range(len(self.artistic.iloc[:, -1]) - 1):
            if ((x[i] == y[i] == 1)):
                w.append("Same boxes of single row were checked")
            elif x[i] == y[i] == 0:
                w.append("Unchecked Boxes")
            else:
                w.append("Completed")
        if w.count("Completed") == len(w):
            self.label_21.setText(u"Ολοκληρώθηκε Επιτυχώς")
            self.chkBoxItem5_5_6.setChecked(True)
        else:
            self.label_21.setText(u"Σφάλμα Συμπλήρωσης")
            self.chkBoxItem5_5_6.setChecked(False)

    def pressed_button22(self):
        a = self.labelResult1_4_19.text()
        b = self.labelResult2_4_19.text()
        c = self.labelResult3_4_19.text()
        d = self.labelResult4_4_19.text()
        e = self.labelResult5_4_19.text()
        f = self.labelResult6_4_19.text()

        # print(a0)
        x0 = a.split(",")
        x0.pop(0)
        x = [1 if i == " True" else 0 for i in x0]
        y0 = b.split(",")
        y0.pop(0)
        y = [1 if i == " True" else 0 for i in y0]

        z0 = c.split(",")
        z0.pop(0)
        z = [1 if i == " True" else 0 for i in z0]
        m0 = d.split(",")
        m0.pop(0)
        m = [1 if i == " True" else 0 for i in m0]

        n0 = e.split(",")
        n0.pop(0)
        n = [1 if i == " True" else 0 for i in n0]
        k0 = f.split(",")
        k0.pop(0)
        k = [1 if i == " True" else 0 for i in k0]

        w1, w2, w3, w4, w5, w6 = [], [], [], [], [], []

        for i in range(len(self.self_assessment__1.iloc[:, -1]) - 1):
            if x[i] == 0:
                w1.append("unchecked")
            else:
                w1.append("check")

            if y[i] == 0:
                w2.append("unchecked")
            else:
                w2.append("check")

            if z[i] == 0:
                w3.append("unchecked")
            else:
                w3.append("check")

            if m[i] == 0:
                w4.append("unchecked")
            else:
                w4.append("check")

            if n[i] == 0:
                w5.append("unchecked")
            else:
                w5.append("check")

            if k[i] == 0:
                w6.append("unchecked")
            else:
                w6.append("check")
                # ~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~
        # ~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~
        # ~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~//~
        a0 = self.labelResult1_4_20.text()
        b0 = self.labelResult2_4_20.text()
        c0 = self.labelResult3_4_20.text()
        d0 = self.labelResult4_4_20.text()
        e0 = self.labelResult5_4_20.text()
        f0 = self.labelResult6_4_20.text()

        x1 = a0.split(",")
        x1.pop(0)
        x2 = [1 if i == " True" else 0 for i in x1]
        y1 = b0.split(",")
        y1.pop(0)
        y2 = [1 if i == " True" else 0 for i in y1]

        z1 = c0.split(",")
        z1.pop(0)
        z2 = [1 if i == " True" else 0 for i in z1]

        m1 = d0.split(",")
        m1.pop(0)
        m2 = [1 if i == " True" else 0 for i in m1]

        n1 = e0.split(",")
        n1.pop(0)
        n2 = [1 if i == " True" else 0 for i in n1]

        k1 = f0.split(",")
        k1.pop(0)
        k2 = [1 if i == " True" else 0 for i in k1]

        v1, v2, v3, v4, v5, v6 = [], [], [], [], [], []

        for i in range(len(self.self_assessment__2.iloc[:, -1]) - 1):
            if x2[i] == 0:
                v1.append("unchecked")
            else:
                v1.append("check")

            if y2[i] == 0:
                v2.append("unchecked")
            else:
                v2.append("check")

            if z2[i] == 0:
                v3.append("unchecked")
            else:
                v3.append("check")

            if m2[i] == 0:
                v4.append("unchecked")
            else:
                v4.append("check")

            if n2[i] == 0:
                v5.append("unchecked")
            else:
                v5.append("check")

            if k2[i] == 0:
                v6.append("unchecked")
            else:
                v6.append("check")

        if ((v1.count("check") == 1) and (v2.count("check") == 1) and (v3.count("check") == 1) and
                (v4.count("check") == 1) and (v5.count("check") == 1) and (v6.count("check") == 1) and
                (w1.count("check") == 1) and (w2.count("check") == 1) and (w3.count("check") == 1) and
                (w4.count("check") == 1) and (w5.count("check") == 1) and (w6.count("check") == 1)):
            self.label_22.setText(u"Ολοκληρώθηκε Επιτυχώς")
        else:
            self.label_22.setText(u"Σφάλμα Συμπλήρωσης")

        if ((v1.count("check") == 1) and (v2.count("check") == 1) and (v3.count("check") == 1) and
                (v4.count("check") == 1) and (v5.count("check") == 1) and (v6.count("check") == 1)):
            self.chkBoxItem6_1_1.setChecked(True)

        else:
            self.chkBoxItem6_1_1.setChecked(False)

        if ((w1.count("check") == 1) and (w2.count("check") == 1) and (w3.count("check") == 1) and
                (w4.count("check") == 1) and (w5.count("check") == 1) and (w6.count("check") == 1)):
            self.chkBoxItem6_1_2.setChecked(True)
        else:
            self.chkBoxItem6_1_2.setChecked(False)

    def word(self):

        self.intr = "Το ερωτηματολόγιο συμπληρώθηκε από τον ενδιαφερόμενο με τα κάτωθι στοιχεία: "
        self.endiaferomenos_1 = "\n" + str(self.label_44.text()) + " " + str(self.lineEdit.text())
        self.endiaferomenos_2 = "\n" + str(self.label_45.text()) + " " + str(self.lineEdit_2.text())
        self.endiaferomenos_3 = "\n" + str(self.label_46.text()) + " " + str(self.comboBox.currentText())
        self.endiaferomenos_4 = "\n" + str(self.label_47.text()) + " " + str(self.comboBox_2.currentText())
        self.endiaferomenos_5 = "\n" + str(self.label_48.text()) + " " + str(self.lineEdit_3.text())

        self.erwtimatologio_aksiwn = '''Το π΄ρωτο ερωτηματολόγιο που συμπλήρωσε ο ενδιαφερόμενος ήταν το Ερωτηματολόγιο Αξιών 
                                        σύμφωνα με το όποίο κατατάσσεται στην κατηγορία των ψυχομετρικών τεστ. Αποτελεί ένα 
                                        εναλλακτικό και λιγότερο πολύπλοκο εργαλείο με σκοπό να μετρήσει τις αξίες των ατόμων 
                                        με βάση ένα μοντέλο 10 θεμελιωδών αξιών, που δημιούργησαν ο Schwartz και οι συνεργάτες
                                        του. Ο Schwartz και οι συνεργάτες του πρότειναν μια ολοκληρωμένη θεωρία μιας καθολικής
                                        ψυχολογικής δομής των ανθρώπινων αξιών. Μέσω του ερωτηματολογίου, αυτή η θεωρία έχει 
                                        σκοπό να εντοπίζει 10 διαφορετικές προτεραιότητες κινήτρων κοινές για ανθρώπους πολλών 
                                        πολιτισμών και κοινωνιών. '''

        self.textEdit_3.setText(self.intr + self.endiaferomenos_1 + self.endiaferomenos_2 + self.endiaferomenos_3
                                + self.endiaferomenos_4 + self.endiaferomenos_5 + "\n" + self.erwtimatologio_aksiwn)


def main():
    app = QtWidgets.QApplication(sys.argv)

    window = MainWindow()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()



